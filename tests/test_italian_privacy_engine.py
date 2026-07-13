from __future__ import annotations

import asyncio
import base64
from io import BytesIO
import json
import tempfile
import subprocess
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader, PdfWriter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from privacy_guardian.activity_log import (
    build_activity_entry,
    export_activity_log_csv,
    load_activity_entries,
    record_activity,
)
from privacy_guardian import document_service
from privacy_guardian.document_service import (
    OcrPageText,
    OcrWord,
    anonymize_loaded_document,
    excluded_value_pairs,
    load_document,
)
from privacy_guardian.models import Finding
from privacy_guardian.privacy_engine import PrivacyEngine
from privacy_guardian.reporting import entity_label, entity_placeholder, report_payload, report_text, source_label
from privacy_guardian.reversible import (
    MAP_SCHEMA_VERSION,
    ReversibleMapEntry,
    ReversibleMapError,
    decrypt_mapping,
    encrypt_mapping,
    read_encrypted_mapping,
    restore_text,
    write_encrypted_mapping,
)
from privacy_guardian.web_app import (
    MAX_TEXT_LENGTH,
    TextPayload,
    analyze_document,
    anonymize as anonymize_text_endpoint,
    anonymize_document,
    engine as web_app_engine,
    health as health_endpoint,
)


class ItalianPrivacyEngineTest(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = PrivacyEngine()

    def findings_for(self, text: str) -> list[tuple[str, str]]:
        return [(finding.entity_type, text[finding.start : finding.end]) for finding in self.engine.analyze(text)]

    def test_does_not_anonymize_common_italian_legal_words(self) -> None:
        self.assertEqual(self.findings_for("Premesso che il contratto viene firmato oggi."), [])
        self.assertEqual(self.findings_for("La società ha premesso quanto segue."), [])

    def test_detects_person_only_with_strong_context(self) -> None:
        findings = self.findings_for("Il sottoscritto Mario Rossi nato a Roma il 10/01/1980.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)
        self.assertNotIn(("DATE_TIME", "10/01/1980"), findings)
        self.assertEqual(self.findings_for("Mario andrà domani a Milano."), [])

    def test_detects_person_with_abbreviated_sig_title(self) -> None:
        findings = self.findings_for("Il Sig. Mario Rossi ha firmato il contratto.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)

        findings = self.findings_for("Gentile sig Mario Rossi, la contattiamo per il rinnovo.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)

        findings = self.findings_for("La Sig.ra Maria Bianchi ha firmato.")
        self.assertIn(("PERSON", "Maria Bianchi"), findings)

    def test_detects_person_when_strong_context_follows_name(self) -> None:
        findings = self.findings_for("Mario Rossi, nato a Roma il 10/01/1980.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)

        findings = self.findings_for("Maria Bianchi residente in Via Appia 12.")
        self.assertIn(("PERSON", "Maria Bianchi"), findings)

        findings = self.findings_for("Bonifico intestato a Mario Rossi.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)

        # Un nome e cognome capitalizzati senza alcun contesto sono comunque un dato
        # personale reale: da quando esiste il dizionario di nomi italiani (vedi
        # DictionaryPersonTest) "Mario Rossi" viene rilevato anche qui, colmando il
        # gap per cui frasi comuni come questa restavano invisibili all'anonimizzatore.
        findings = self.findings_for("Mario Rossi andrà domani a Milano.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)
        self.assertNotIn(("PERSON", "Milano"), findings)

    def test_does_not_anonymize_dates(self) -> None:
        text = "Il sottoscritto Mario Rossi nato il 10/01/1980 e residente dal 5 maggio 2020."
        anonymized = self.engine.anonymize(text)
        self.assertIn("10/01/1980", anonymized)
        self.assertIn("5 maggio 2020", anonymized)

    def test_maximum_protection_redacts_dates_and_initials(self) -> None:
        text = (
            "Il sottoscritto Mario Rossi nato il 10/01/1980 lavora per "
            "Alfa Beta S.r.l. in Via Appia 12."
        )
        anonymized = self.engine.anonymize(text, mode="maximum")

        self.assertIn("<PERSONA>", anonymized)
        self.assertIn("<DATA>", anonymized)
        self.assertIn("<ORGANIZZAZIONE>", anonymized)
        self.assertIn("<INDIRIZZO>", anonymized)
        self.assertNotIn("M. R.", anonymized)
        self.assertNotIn("10/01/1980", anonymized)

    def test_reversible_mode_uses_numbered_placeholders_and_restores_text(self) -> None:
        text = (
            "Il sottoscritto Mario Rossi email mario@example.com nato il 10/01/1980. "
            "Bonifico intestato a Mario Rossi."
        )
        findings = self.engine.analyze(text, mode="reversible")
        result = self.engine.anonymize_reversible(text, findings)

        self.assertEqual(result.text.count("<PERSONA_1>"), 2)
        self.assertIn("<EMAIL_1>", result.text)
        self.assertIn("<DATA_1>", result.text)
        self.assertNotIn("Mario Rossi", result.text)
        self.assertNotIn("mario@example.com", result.text)
        self.assertEqual(restore_text(result.text, result.mapping), text)

        encrypted = encrypt_mapping(result.mapping, "password locale")
        self.assertNotIn(b"Mario Rossi", encrypted)
        self.assertNotIn(b"mario@example.com", encrypted)
        decrypted_mapping = decrypt_mapping(encrypted, "password locale")
        self.assertEqual(restore_text(result.text, decrypted_mapping), text)

    def test_reversible_placeholders_skip_ones_already_present_in_text(self) -> None:
        text = "Nota: usa <PERSONA_1> come esempio. Il sottoscritto Mario Rossi conferma."
        findings = self.engine.analyze(text, mode="reversible")
        result = self.engine.anonymize_reversible(text, findings)

        self.assertIn("<PERSONA_1> come esempio", result.text)
        self.assertIn("<PERSONA_2>", result.text)
        self.assertNotIn("Mario Rossi", result.text)
        self.assertEqual(restore_text(result.text, result.mapping), text)

    def test_reuses_preloaded_reversible_map_across_documents(self) -> None:
        text = "Il sottoscritto Mario Rossi email mario@example.com conferma."
        preloaded = (ReversibleMapEntry(placeholder="<PERSONA_1>", entity_type="PERSON", value="Mario Rossi"),)

        findings = self.engine.analyze(text, mode="reversible")
        result = self.engine.anonymize_reversible(text, findings, entries=preloaded)

        self.assertIn("<PERSONA_1>", result.text)
        self.assertNotIn("<PERSONA_2>", result.text)
        self.assertIn("<EMAIL_1>", result.text)
        self.assertIn(preloaded[0], result.mapping)
        self.assertTrue(
            any(entry.entity_type == "EMAIL_ADDRESS" and entry.value == "mario@example.com" for entry in result.mapping)
        )
        self.assertEqual(restore_text(result.text, result.mapping), text)

    def test_dates_are_findings_only_in_maximum_protection(self) -> None:
        text = "Il sottoscritto Mario Rossi nato il 5 maggio 2020."

        standard_findings = self.findings_for(text)
        maximum_findings = [
            (finding.entity_type, text[finding.start : finding.end])
            for finding in self.engine.analyze(text, mode="maximum")
        ]

        self.assertNotIn(("DATE", "5 maggio 2020"), standard_findings)
        self.assertIn(("DATE", "5 maggio 2020"), maximum_findings)

    def test_detects_italian_organizations(self) -> None:
        findings = self.findings_for("Alfa Beta S.r.l. invierà la fattura.")
        self.assertIn(("ORGANIZATION", "Alfa Beta S.r.l."), findings)

        findings = self.findings_for("società Rossi Consulting spa con sede in Via Garibaldi 12, Milano.")
        self.assertIn(("ORGANIZATION", "società Rossi Consulting spa"), findings)
        self.assertIn(("ADDRESS", "Via Garibaldi 12, Milano"), findings)

        findings = self.findings_for("ditta individuale Mario Rossi P.IVA 12345678903")
        self.assertIn(("ORGANIZATION", "ditta individuale Mario Rossi"), findings)
        self.assertIn(("PARTITA_IVA", "12345678903"), findings)

    def test_detects_organizations_with_ampersand(self) -> None:
        findings = self.findings_for("Contratto con Rossi & Figli S.r.l. per la fornitura.")
        self.assertIn(("ORGANIZATION", "Rossi & Figli S.r.l."), findings)

        anonymized = self.engine.anonymize("La Rossi & Figli S.r.l. ha sede a Milano.", mode="maximum")
        self.assertNotIn("Rossi", anonymized)
        self.assertNotIn("Figli", anonymized)
        self.assertIn("<ORGANIZZAZIONE>", anonymized)

    def test_keeps_initials_for_people_organizations_and_places(self) -> None:
        text = "Il sottoscritto Mario Rossi lavora per Alfa Beta S.r.l. nella Provincia di Potenza."
        anonymized = self.engine.anonymize(text)
        self.assertIn("M. R.", anonymized)
        self.assertIn("A. B. S. r. l.", anonymized)
        self.assertIn("P. d. P.", anonymized)
        self.assertNotIn("<PERSONA>", anonymized)
        self.assertNotIn("<ORGANIZZAZIONE>", anonymized)

    def test_detects_territorial_bodies(self) -> None:
        findings = self.findings_for("Provincia di Potenza, regione Basilicata e Comune di Roma.")
        self.assertIn(("TERRITORIAL_BODY", "Provincia di Potenza"), findings)
        self.assertIn(("TERRITORIAL_BODY", "regione Basilicata"), findings)
        self.assertIn(("TERRITORIAL_BODY", "Comune di Roma"), findings)

    def test_detects_qualified_administration_with_place(self) -> None:
        findings = self.findings_for("amministrazione provinciale di Potenza")
        self.assertIn(("TERRITORIAL_BODY", "amministrazione provinciale di Potenza"), findings)

        anonymized = self.engine.anonymize("amministrazione provinciale di Potenza")
        self.assertNotIn("Potenza", anonymized)

    def test_detects_capitalized_qualified_administration_with_place(self) -> None:
        findings = self.findings_for("Amministrazione Provinciale di Potenza")
        self.assertIn(("TERRITORIAL_BODY", "Amministrazione Provinciale di Potenza"), findings)

    def test_unqualified_administration_is_not_a_territorial_body(self) -> None:
        # "amministrazione del personale" non è un ente territoriale: senza il
        # qualificatore provinciale/comunale/regionale non deve scattare il falso positivo.
        findings = self.findings_for("amministrazione del personale")
        self.assertFalse(any(entity_type == "TERRITORIAL_BODY" for entity_type, _ in findings))

    def test_detects_prefettura_and_questura_with_place(self) -> None:
        findings = self.findings_for("Prefettura di Matera")
        self.assertIn(("TERRITORIAL_BODY", "Prefettura di Matera"), findings)

        findings = self.findings_for("Questura di Bari")
        self.assertIn(("TERRITORIAL_BODY", "Questura di Bari"), findings)

    def test_detects_addresses_without_house_number(self) -> None:
        findings = self.findings_for("Residente in Via Appia e domiciliato in Viale Europa 10.")
        self.assertIn(("ADDRESS", "Via Appia"), findings)
        self.assertIn(("ADDRESS", "Viale Europa 10"), findings)

    def test_full_address_with_cap_and_city_is_a_single_finding(self) -> None:
        text = "Via Garibaldi 45, 00185 Roma"
        findings = self.engine.analyze(text, "maximum")
        addresses = [finding for finding in findings if finding.entity_type == "ADDRESS"]

        self.assertEqual(len(addresses), 1)
        self.assertEqual(text[addresses[0].start : addresses[0].end], text)

        anonymized = self.engine.anonymize(text, findings, "maximum")
        self.assertNotIn("Roma", anonymized)
        self.assertNotIn("00185", anonymized)
        self.assertFalse(any(char.isdigit() for char in anonymized))

    def test_full_address_without_comma_is_fully_covered(self) -> None:
        text = "Via Garibaldi 45 00185 Roma"
        findings = self.engine.analyze(text, "maximum")
        addresses = [finding for finding in findings if finding.entity_type == "ADDRESS"]

        self.assertEqual(len(addresses), 1)
        self.assertEqual(text[addresses[0].start : addresses[0].end], text)

        anonymized = self.engine.anonymize(text, findings, "maximum")
        self.assertNotIn("Roma", anonymized)
        self.assertNotIn("00185", anonymized)
        self.assertFalse(any(char.isdigit() for char in anonymized))

    def test_full_address_with_alphanumeric_civic_is_fully_covered(self) -> None:
        text = "Via Roma 45B, 00185 Roma"
        findings = self.engine.analyze(text, "maximum")
        addresses = [finding for finding in findings if finding.entity_type == "ADDRESS"]

        self.assertEqual(len(addresses), 1)
        self.assertEqual(text[addresses[0].start : addresses[0].end], text)

        anonymized = self.engine.anonymize(text, findings, "maximum")
        self.assertNotIn("Roma", anonymized)
        self.assertNotIn("00185", anonymized)
        self.assertFalse(any(char.isdigit() for char in anonymized))

    def test_detects_structured_italian_data(self) -> None:
        text = "IBAN IT60X0542811101000000123456 CF RSSMRA80A01H501U email mario.rossi@example.com tel +39 333 123 4567"
        anonymized = self.engine.anonymize(text)
        self.assertIn("<IBAN>", anonymized)
        self.assertIn("<CODICE_FISCALE>", anonymized)
        self.assertIn("<EMAIL>", anonymized)
        self.assertIn("<TELEFONO>", anonymized)

    def test_detects_labeled_spaced_codice_fiscale(self) -> None:
        text = "C.F. RSS MRA 80A01 H501U"
        findings = self.findings_for(text)
        self.assertIn(("CODICE_FISCALE", "RSS MRA 80A01 H501U"), findings)

        anonymized = self.engine.anonymize(text)
        self.assertNotIn("RSS MRA 80A01 H501U", anonymized)
        self.assertIn("<CODICE_FISCALE>", anonymized)

    def test_labeled_contiguous_codice_fiscale_still_detected(self) -> None:
        findings = self.findings_for("CF: RSSMRA80A01H501U")
        self.assertIn(("CODICE_FISCALE", "RSSMRA80A01H501U"), findings)

    def test_labeled_codice_fiscale_detected_even_with_invalid_checksum(self) -> None:
        # Struttura plausibile a 16 caratteri ma checksum errato: l'etichetta è un
        # contesto abbastanza forte da trattarlo comunque come codice fiscale.
        text = "codice fiscale RSSMRA80A01H501Z"
        self.assertFalse(self.engine._recognizer._valid_codice_fiscale("RSSMRA80A01H501Z"))
        findings = self.findings_for(text)
        self.assertIn(("CODICE_FISCALE", "RSSMRA80A01H501Z"), findings)

    def test_labeled_codice_fiscale_requires_uppercase_value(self) -> None:
        # Il requisito è esplicito: dopo l'etichetta il valore deve essere in maiuscolo,
        # anche se l'etichetta stessa tollera varianti di maiuscole/minuscole. Il pattern
        # etichettato/spaziato (con checksum valido) non deve scattare su un valore minuscolo;
        # il caso contiguo minuscolo resta comunque coperto dal riconoscitore CF classico
        # (case-insensitive con validazione del checksum), che è un percorso indipendente.
        recognizer = self.engine._recognizer
        self.assertEqual(recognizer._labeled_codice_fiscale_findings("cf rssmra80a01h501u"), [])
        self.assertEqual(recognizer._labeled_codice_fiscale_findings("cf rss mra 80a01 h501u"), [])

    def test_labeled_codice_fiscale_does_not_overmatch_short_candidates(self) -> None:
        self.assertFalse(any(entity_type == "CODICE_FISCALE" for entity_type, _ in self.findings_for("CF ABC")))
        self.assertFalse(
            any(entity_type == "CODICE_FISCALE" for entity_type, _ in self.findings_for("C.F.R. Milano"))
        )

    def test_detects_foreign_ibans_with_checksum(self) -> None:
        text = "Bonifico su IBAN DE89 3704 0044 0532 0130 00 e FR1420041010050500013M02606 entro oggi."
        findings = self.findings_for(text)

        self.assertIn(("IBAN", "DE89 3704 0044 0532 0130 00"), findings)
        self.assertIn(("IBAN", "FR1420041010050500013M02606"), findings)

        anonymized = self.engine.anonymize(text)
        self.assertEqual(anonymized.count("<IBAN>"), 2)
        self.assertIn("entro oggi", anonymized)

    def test_rejects_iban_candidates_with_wrong_checksum(self) -> None:
        findings = self.findings_for("Conto DE89 3704 0044 0532 0130 01 non valido.")
        self.assertNotIn("IBAN", [entity_type for entity_type, _ in findings])

    def test_detects_credit_card_numbers_in_multiple_formats(self) -> None:
        findings = self.findings_for("La carta è 4111 1111 1111 1111 per il pagamento.")
        self.assertIn(("CREDIT_CARD", "4111 1111 1111 1111"), findings)

        findings = self.findings_for("La carta è 4111-1111-1111-1111 per il pagamento.")
        self.assertIn(("CREDIT_CARD", "4111-1111-1111-1111"), findings)

        findings = self.findings_for("La carta è 4111111111111111 per il pagamento.")
        self.assertIn(("CREDIT_CARD", "4111111111111111"), findings)

    def test_rejects_credit_card_candidates_with_invalid_luhn_checksum(self) -> None:
        findings = self.findings_for("Numero non valido 4111 1111 1111 1112.")
        self.assertNotIn("CREDIT_CARD", [entity_type for entity_type, _ in findings])

    def test_credit_card_digits_inside_iban_remain_iban(self) -> None:
        text = "IBAN IT60X0542811101000000123456 per il bonifico."
        findings = self.findings_for(text)

        self.assertIn(("IBAN", "IT60X0542811101000000123456"), findings)
        self.assertNotIn("CREDIT_CARD", [entity_type for entity_type, _ in findings])

        text = "Bonifico su IBAN DE89 3704 0044 0532 0130 00 e altro."
        findings = self.findings_for(text)
        self.assertIn(("IBAN", "DE89 3704 0044 0532 0130 00"), findings)
        self.assertNotIn("CREDIT_CARD", [entity_type for entity_type, _ in findings])

    def test_maximum_protection_redacts_credit_card_number(self) -> None:
        text = "La carta è 4111 1111 1111 1111 per il pagamento."
        anonymized = self.engine.anonymize(text, mode="maximum")
        self.assertIn("<CARTA_PAGAMENTO>", anonymized)
        self.assertNotIn("4111", anonymized)

    def test_detects_postal_code_with_city_as_address(self) -> None:
        findings = self.findings_for("Abito a 00185 Roma da anni.")
        self.assertIn(("ADDRESS", "00185 Roma"), findings)

        findings = self.findings_for("Sede legale 20121 Milano centro.")
        self.assertIn(("ADDRESS", "20121 Milano"), findings)

        findings = self.findings_for("Ufficio 47521 Cesena aperto.")
        self.assertIn(("ADDRESS", "47521 Cesena"), findings)

    def test_postal_code_without_city_is_not_matched(self) -> None:
        findings = self.findings_for("Il numero 00185 non è una città.")
        self.assertNotIn("ADDRESS", [entity_type for entity_type, _ in findings])

    def test_postal_code_followed_by_month_is_not_matched(self) -> None:
        findings = self.findings_for("Consegna prevista 12345 Gennaio del prossimo anno.")
        self.assertEqual(findings, [])

    def test_amount_followed_by_currency_is_not_matched_as_address(self) -> None:
        findings = self.findings_for("Il compenso pattuito è di 10000 Euro da corrispondere in due rate.")
        self.assertNotIn("ADDRESS", [entity_type for entity_type, _ in findings])

    def test_detects_international_phone_numbers(self) -> None:
        text = "Contatti: +44 20 7946 0958, +1 415 555 0132 e +39 333 123 4567."
        findings = self.findings_for(text)

        self.assertIn(("PHONE_NUMBER", "+44 20 7946 0958"), findings)
        self.assertIn(("PHONE_NUMBER", "+1 415 555 0132"), findings)
        self.assertIn(("PHONE_NUMBER", "+39 333 123 4567"), findings)
        self.assertEqual(self.engine.anonymize(text).count("<TELEFONO>"), 3)

    def test_detects_lowercase_addresses_with_house_number(self) -> None:
        findings = self.findings_for("abito in via giuseppe garibaldi 12, roma")
        self.assertIn(("ADDRESS", "via giuseppe garibaldi 12"), findings)

        findings = self.findings_for("residenza in piazza dei martiri, 4")
        self.assertIn(("ADDRESS", "piazza dei martiri, 4"), findings)

    def test_lowercase_address_idioms_are_not_matched(self) -> None:
        self.assertEqual(self.findings_for("procediamo in via preliminare entro il 12 del mese"), [])
        self.assertEqual(self.findings_for("rispondere via email entro il 15"), [])

    def test_person_name_does_not_swallow_street_address(self) -> None:
        text = "Il dott. Mario Rossi Via Appia 12 chiede accesso."
        findings = self.findings_for(text)

        self.assertIn(("PERSON", "Mario Rossi"), findings)
        address_values = [value for entity_type, value in findings if entity_type == "ADDRESS"]
        self.assertTrue(any(value.startswith("Via Appia 12") for value in address_values))

    def test_propagates_person_coreferences_to_full_name_and_surname(self) -> None:
        text = (
            "Il sig. Mario Rossi ha firmato. Successivamente Rossi ha inviato una lettera "
            "e Mario Rossi ha confermato."
        )
        findings = [
            (finding.entity_type, text[finding.start : finding.end], finding.source)
            for finding in self.engine.analyze(text, "reversible")
        ]

        self.assertIn(("PERSON", "Mario Rossi", "italian_rules"), findings)
        self.assertIn(("PERSON", "Rossi", "coreference"), findings)
        self.assertEqual(sum(1 for entity, value, _ in findings if entity == "PERSON" and value == "Mario Rossi"), 2)
        self.assertEqual(source_label("coreference"), "Propagazione nome")

        result = self.engine.anonymize_reversible(text, self.engine.analyze(text, "reversible"))
        self.assertEqual(result.text.count("<PERSONA_1>"), 2)
        self.assertNotIn("Mario Rossi", result.text)
        self.assertNotIn("Rossi", result.text)
        self.assertEqual(restore_text(result.text, result.mapping), text)

    def test_does_not_propagate_surname_that_is_part_of_an_address(self) -> None:
        text = "Il sig. Mario Verdi abita in via Verdi 3."
        findings = self.findings_for(text)

        self.assertIn(("PERSON", "Mario Verdi"), findings)
        self.assertTrue(any(entity == "ADDRESS" and value.startswith("via Verdi 3") for entity, value in findings))
        self.assertNotIn(("PERSON", "Verdi"), findings)

        anonymized = self.engine.anonymize(text, mode="maximum")
        self.assertIn("<INDIRIZZO>", anonymized)
        self.assertNotIn("Verdi 3", anonymized)

    def test_does_not_propagate_lowercase_surname_occurrences(self) -> None:
        text = "Il sig. Mario Rossi ha firmato. Successivamente rossi ha inviato una lettera."
        findings = self.findings_for(text)

        self.assertIn(("PERSON", "Mario Rossi"), findings)
        self.assertEqual([value for entity, value in findings if entity == "PERSON"], ["Mario Rossi"])
        self.assertIn("rossi", self.engine.anonymize(text, mode="maximum"))

    def test_ner_active_reflects_availability_of_local_ner(self) -> None:
        from types import SimpleNamespace

        from privacy_guardian.ner_recognizer import NerPersonRecognizer

        def fake_nlp(value: str):
            return SimpleNamespace(ents=[])

        with patch(
            "privacy_guardian.privacy_engine.NerPersonRecognizer.create_if_available",
            return_value=NerPersonRecognizer(fake_nlp),
        ):
            engine_with_ner = PrivacyEngine()
        self.assertIsInstance(engine_with_ner.ner_active, bool)
        self.assertTrue(engine_with_ner.ner_active)

        with patch(
            "privacy_guardian.privacy_engine.NerPersonRecognizer.create_if_available",
            return_value=None,
        ):
            engine_without_ner = PrivacyEngine()
        self.assertIsInstance(engine_without_ner.ner_active, bool)
        self.assertFalse(engine_without_ner.ner_active)

        default_engine = PrivacyEngine()
        self.assertIsInstance(default_engine.ner_active, bool)
        self.assertEqual(default_engine.ner_active, default_engine._ner is not None)

    def test_optional_local_ner_adds_uncontexted_person_names(self) -> None:
        from types import SimpleNamespace

        from privacy_guardian.ner_recognizer import NerPersonRecognizer

        # Nomi stranieri, non presenti nel dizionario locale di nomi italiani
        # (first_names.py): questo isola davvero il contributo del NER opzionale,
        # che non dipende dal dizionario. Con nomi italiani comuni come "Mario
        # Rossi" il dizionario locale li rileva ora da solo (score più alto),
        # vedi DictionaryPersonTest.
        text = "Wolfgang Keller ha inviato la relazione a Elin Andersson."

        def fake_nlp(value: str):
            ents = []
            for name in ("Wolfgang Keller", "Elin Andersson"):
                start = value.find(name)
                if start >= 0:
                    ents.append(
                        SimpleNamespace(label_="PER", text=name, start_char=start, end_char=start + len(name))
                    )
            return SimpleNamespace(ents=ents)

        engine = PrivacyEngine()
        engine._ner = NerPersonRecognizer(fake_nlp)

        findings = [
            (finding.entity_type, text[finding.start : finding.end], finding.source)
            for finding in engine.analyze(text, "maximum")
        ]
        self.assertIn(("PERSON", "Wolfgang Keller", "ner_local"), findings)
        self.assertIn(("PERSON", "Elin Andersson", "ner_local"), findings)

        anonymized = engine.anonymize(text, mode="maximum")
        self.assertNotIn("Wolfgang Keller", anonymized)
        self.assertNotIn("Elin Andersson", anonymized)
        self.assertEqual(source_label("ner_local"), "NER locale (spaCy)")

    def test_detects_pec_as_dedicated_category(self) -> None:
        text = (
            "Email ordinaria mario.rossi@example.com, PEC azienda@pec.it "
            "e domicilio digitale studio.rossi@example.com."
        )
        findings = self.findings_for(text)
        anonymized = self.engine.anonymize(text)

        self.assertIn(("EMAIL_ADDRESS", "mario.rossi@example.com"), findings)
        self.assertIn(("PEC_ADDRESS", "azienda@pec.it"), findings)
        self.assertIn(("PEC_ADDRESS", "studio.rossi@example.com"), findings)
        self.assertNotIn(("EMAIL_ADDRESS", "azienda@pec.it"), findings)
        self.assertEqual(anonymized.count("<EMAIL>"), 1)
        self.assertEqual(anonymized.count("<PEC>"), 2)

    def test_detects_common_italian_number_formats(self) -> None:
        text = "IBAN IT60 X054 2811 1010 0000 0123 456 tel 06/12345678 cell +39 333/1234567"
        findings = self.findings_for(text)
        anonymized = self.engine.anonymize(text)

        self.assertIn(("IBAN", "IT60 X054 2811 1010 0000 0123 456"), findings)
        self.assertIn(("PHONE_NUMBER", "06/12345678"), findings)
        self.assertIn(("PHONE_NUMBER", "+39 333/1234567"), findings)
        self.assertIn("<IBAN>", anonymized)
        self.assertEqual(anonymized.count("<TELEFONO>"), 2)
        self.assertNotIn("IT60 X054", anonymized)
        self.assertNotIn("0123 456", anonymized)

    def test_detects_invoice_and_health_identifiers_with_context(self) -> None:
        text = (
            "Codice destinatario SDI ABC1234, codice univoco ufficio A1B2C3 "
            "e tessera sanitaria n. 8038 0000 0000 0000 0000."
        )
        findings = self.findings_for(text)
        anonymized = self.engine.anonymize(text)

        self.assertIn(("SDI_CODE", "ABC1234"), findings)
        self.assertIn(("SDI_CODE", "A1B2C3"), findings)
        self.assertIn(("HEALTH_CARD", "8038 0000 0000 0000 0000"), findings)
        self.assertEqual(anonymized.count("<CODICE_SDI>"), 2)
        self.assertIn("<TESSERA_SANITARIA>", anonymized)
        self.assertNotIn("ABC1234", anonymized)
        self.assertNotIn("8038 0000", anonymized)

    def test_detects_protocol_and_case_numbers_with_context(self) -> None:
        text = (
            "Protocollo n. 12345/2024, prot. SUAP-7788-A, "
            "pratica ED 9901 e fascicolo RG 4567/2023."
        )
        findings = self.findings_for(text)
        anonymized = self.engine.anonymize(text)

        self.assertIn(("PROTOCOL_CASE_NUMBER", "12345/2024"), findings)
        self.assertIn(("PROTOCOL_CASE_NUMBER", "SUAP-7788-A"), findings)
        self.assertIn(("PROTOCOL_CASE_NUMBER", "ED 9901"), findings)
        self.assertIn(("PROTOCOL_CASE_NUMBER", "RG 4567/2023"), findings)
        self.assertEqual(anonymized.count("<PROTOCOLLO_PRATICA>"), 4)

    def test_protocol_and_case_numbers_require_context_and_skip_dates(self) -> None:
        text = "La sigla SUAP-7788-A non basta. Protocollo del 10/01/2024 senza numero pratica."
        findings = self.findings_for(text)

        self.assertNotIn(("PROTOCOL_CASE_NUMBER", "SUAP-7788-A"), findings)
        self.assertNotIn(("PROTOCOL_CASE_NUMBER", "10/01/2024"), findings)

    def test_invoice_and_health_codes_require_clear_context(self) -> None:
        text = "La sigla ABC1234 e il numero 80380000000000000000 non bastano da soli."

        self.assertNotIn(("SDI_CODE", "ABC1234"), self.findings_for(text))
        self.assertNotIn(("HEALTH_CARD", "80380000000000000000"), self.findings_for(text))

    def test_detects_identity_documents_and_vehicle_plates_with_context(self) -> None:
        text = (
            "Carta d'identità n. CA12345AA, passaporto n. YA1234567, "
            "patente n. U1234567A e targa AB123CD."
        )
        findings = self.findings_for(text)
        anonymized = self.engine.anonymize(text)

        self.assertIn(("IDENTITY_DOCUMENT", "CA12345AA"), findings)
        self.assertIn(("IDENTITY_DOCUMENT", "YA1234567"), findings)
        self.assertIn(("IDENTITY_DOCUMENT", "U1234567A"), findings)
        self.assertIn(("VEHICLE_PLATE", "AB123CD"), findings)
        self.assertEqual(anonymized.count("<DOCUMENTO_IDENTITA>"), 3)
        self.assertIn("<TARGA_VEICOLO>", anonymized)
        self.assertNotIn("CA12345AA", anonymized)
        self.assertNotIn("AB123CD", anonymized)

    def test_document_and_plate_codes_require_clear_context(self) -> None:
        text = "La pratica CA12345AA e la sigla AB123CD non sono sufficienti da sole."

        self.assertNotIn(("IDENTITY_DOCUMENT", "CA12345AA"), self.findings_for(text))
        self.assertNotIn(("VEHICLE_PLATE", "AB123CD"), self.findings_for(text))

    def test_report_summarizes_mode_counts_and_review_warning(self) -> None:
        text = "Il sottoscritto Mario Rossi email mario@example.com nato il 10/01/1980."
        findings = self.engine.analyze(text, mode="maximum")
        report = report_payload(findings, "maximum")

        self.assertEqual(report["mode_label"], "Massima protezione")
        self.assertEqual(report["counts"]["PERSON"], 1)
        self.assertEqual(report["counts"]["EMAIL_ADDRESS"], 1)
        self.assertEqual(report["counts"]["DATE"], 1)
        self.assertIn("Rileggi sempre", report["summary"])
        self.assertIn("checklist", report)
        self.assertIn("Massima protezione", report["checklist"][0])
        self.assertIn("1 persona", report["summary"])
        self.assertIn("1 data", report["summary"])
        self.assertIn("3 dati riconosciuti", report_text(findings, "maximum"))

    def test_standard_report_warns_about_initials_and_dates(self) -> None:
        report = report_text([], "standard")

        self.assertIn("Standard conserva iniziali e date", report)
        self.assertIn("0 dati riconosciuti", report)

        payload = report_payload([], "standard")
        self.assertIn("Standard lascia visibili iniziali e date", payload["checklist"][0])

    def test_reversible_report_explains_local_map(self) -> None:
        payload = report_payload([], "reversible")

        self.assertEqual(payload["mode_label"], "Reversibile")
        self.assertIn("mappa locale cifrata", payload["mode_note"])
        self.assertIn("salva la mappa cifrata", payload["checklist"][0])

    def test_entity_labels_are_human_readable(self) -> None:
        self.assertEqual(entity_label("PERSON"), "persona")
        self.assertEqual(entity_label("CREDIT_CARD"), "carta di pagamento")
        self.assertEqual(entity_label("CREDIT_CARD", 2), "carte di pagamento")
        self.assertEqual(entity_placeholder("CREDIT_CARD"), "<CARTA_PAGAMENTO>")
        self.assertEqual(entity_label("PHONE_NUMBER", 2), "telefoni")
        self.assertEqual(entity_label("IDENTITY_DOCUMENT"), "documento d'identità")
        self.assertEqual(entity_label("VEHICLE_PLATE", 2), "targhe veicolo")
        self.assertEqual(entity_label("HEALTH_CARD"), "tessera sanitaria")
        self.assertEqual(entity_label("SDI_CODE", 2), "codici SDI")
        self.assertEqual(entity_label("PEC_ADDRESS"), "PEC")
        self.assertEqual(entity_label("PROTOCOL_CASE_NUMBER", 2), "numeri protocollo/pratica")
        self.assertEqual(entity_placeholder("PERSON"), "<PERSONA>")
        self.assertEqual(entity_placeholder("IDENTITY_DOCUMENT"), "<DOCUMENTO_IDENTITA>")
        self.assertEqual(entity_placeholder("PHONE_NUMBER"), "<TELEFONO>")
        self.assertEqual(entity_placeholder("HEALTH_CARD"), "<TESSERA_SANITARIA>")
        self.assertEqual(entity_placeholder("SDI_CODE"), "<CODICE_SDI>")
        self.assertEqual(entity_placeholder("PEC_ADDRESS"), "<PEC>")
        self.assertEqual(entity_placeholder("PROTOCOL_CASE_NUMBER"), "<PROTOCOLLO_PRATICA>")
        self.assertEqual(source_label("italian_rules"), "Regole italiane locali")


class DictionaryPersonTest(unittest.TestCase):
    """Riconoscimento di nomi di persona senza titoli né contesto forte, tramite il
    dizionario locale di nomi italiani (src/privacy_guardian/assets/nomi_italiani.txt)."""

    def setUp(self) -> None:
        self.engine = PrivacyEngine()

    def findings_for(self, text: str) -> list[tuple[str, str]]:
        return [(finding.entity_type, text[finding.start : finding.end]) for finding in self.engine.analyze(text)]

    def test_detects_person_in_common_phrase_without_titles(self) -> None:
        text = "Vi scrivo in merito alla pratica di Mario Rossi, aperta il mese scorso."
        findings = self.findings_for(text)
        self.assertIn(("PERSON", "Mario Rossi"), findings)

        anonymized = self.engine.anonymize(text)
        self.assertNotIn("Mario Rossi", anonymized)

    def test_detects_person_at_start_of_sentence(self) -> None:
        findings = self.findings_for("Mario Rossi ha richiesto copia del contratto.")
        self.assertIn(("PERSON", "Mario Rossi"), findings)

    def test_detects_person_after_generic_verb(self) -> None:
        findings = self.findings_for("Ho parlato con Giulia Bianchi riguardo al preventivo.")
        self.assertIn(("PERSON", "Giulia Bianchi"), findings)

    def test_propagates_coreference_from_dictionary_match(self) -> None:
        text = "La pratica di Mario Rossi è aperta. Rossi ha chiesto copia."
        findings = [
            (finding.entity_type, text[finding.start : finding.end], finding.source)
            for finding in self.engine.analyze(text)
        ]
        self.assertIn(("PERSON", "Mario Rossi", "name_dictionary"), findings)
        self.assertIn(("PERSON", "Rossi", "coreference"), findings)
        self.assertEqual(sum(1 for entity, _, _ in findings if entity == "PERSON"), 2)

    def test_detects_person_preceded_by_capitalized_non_name_word(self) -> None:
        # Una maiuscola non-nome prima del nome ("Repubblica", "Milano") non deve
        # nascondere la persona: la scansione è ancorata al nome di battesimo.
        for text in (
            "Il Presidente della Repubblica Sergio Mattarella ha firmato.",
            "Da Milano Sergio Mattarella è partito ieri.",
        ):
            findings = self.findings_for(text)
            self.assertIn(("PERSON", "Sergio Mattarella"), findings, text)

    def test_street_named_after_a_person_is_not_flagged_as_person(self) -> None:
        text = "L'ufficio si trova in via Mario Rossi 12, 00185 Roma."
        findings = self.findings_for(text)

        self.assertNotIn(("PERSON", "Mario Rossi"), findings)
        self.assertFalse(any(entity_type == "PERSON" for entity_type, _ in findings))

    def test_lowercase_name_is_not_matched(self) -> None:
        self.assertEqual(self.findings_for("la pratica di mario rossi"), [])

    def test_capitalized_non_name_word_is_not_matched(self) -> None:
        # "Milano" non è un nome di battesimo: solo "Comune di Milano Marittima"
        # come TERRITORIAL_BODY, nessun PERSON.
        findings = self.findings_for("Il Comune di Milano Marittima")
        self.assertFalse(any(entity_type == "PERSON" for entity_type, _ in findings))

    def test_stopword_prefix_is_still_handled_as_before(self) -> None:
        # "Cliente Rossi": "Cliente" non è un nome nel dizionario e "Rossi" da solo
        # non forma un candidato di due parole, quindi nessuna regressione.
        self.assertEqual(self.findings_for("Cliente Rossi ha firmato."), [])

    def test_dictionary_loads_and_contains_expected_names(self) -> None:
        from privacy_guardian.first_names import is_italian_first_name, load_italian_first_names

        names = load_italian_first_names()
        self.assertIsInstance(names, frozenset)
        self.assertGreater(len(names), 0)
        for expected in ("mario", "giulia", "francesca"):
            self.assertIn(expected, names)
        for unexpected in ("milano", "roma", "via"):
            self.assertNotIn(unexpected, names)

        self.assertTrue(is_italian_first_name("Mario"))
        self.assertTrue(is_italian_first_name("GIULIA,"))
        self.assertFalse(is_italian_first_name("Milano"))
        self.assertFalse(is_italian_first_name(""))


class ExcludedValuePairsTest(unittest.TestCase):
    """Fail-closed computation of the (entity_type, value) pairs excluded by the UI mask."""

    def test_pair_excluded_only_when_every_occurrence_is_unchecked(self) -> None:
        email = "mario@example.com"
        text = f"{email} conferma da {email}"
        second = text.index(email, 1)
        findings = [
            Finding("EMAIL_ADDRESS", 0, len(email), 0.98),
            Finding("EMAIL_ADDRESS", second, second + len(email), 0.98),
        ]

        self.assertEqual(excluded_value_pairs(text, findings, [False, True]), frozenset())
        self.assertEqual(excluded_value_pairs(text, findings, [True, False]), frozenset())
        self.assertEqual(
            excluded_value_pairs(text, findings, [False, False]),
            frozenset({("EMAIL_ADDRESS", email)}),
        )

    def test_short_mask_counts_missing_indices_as_included(self) -> None:
        email = "mario@example.com"
        text = f"{email} conferma da {email}"
        second = text.index(email, 1)
        findings = [
            Finding("EMAIL_ADDRESS", 0, len(email), 0.98),
            Finding("EMAIL_ADDRESS", second, second + len(email), 0.98),
        ]

        self.assertEqual(excluded_value_pairs(text, findings, [False]), frozenset())
        self.assertEqual(excluded_value_pairs(text, findings, []), frozenset())


class AddExtraValueFindingsTest(unittest.TestCase):
    """add_extra_value_findings aggiunge un finding per ogni occorrenza letterale di un
    valore selezionato manualmente, senza duplicare quelle già coperte."""

    def test_adds_a_finding_for_every_occurrence(self) -> None:
        value = "ABC9988"
        text = f"Documento {value} rilasciato, copia {value} allegata."
        result = document_service.add_extra_value_findings(text, [], frozenset({("IDENTITY_DOCUMENT", value)}))

        self.assertEqual(len(result), 2)
        for finding in result:
            self.assertEqual(finding.entity_type, "IDENTITY_DOCUMENT")
            self.assertEqual(text[finding.start : finding.end], value)
            self.assertEqual(finding.source, "manual")

    def test_skips_occurrence_overlapping_an_existing_finding(self) -> None:
        email = "mario@example.com"
        text = f"Contatto: {email}"
        existing = Finding("EMAIL_ADDRESS", text.index(email), text.index(email) + len(email), 0.98)

        result = document_service.add_extra_value_findings(text, [existing], frozenset({("EMAIL_ADDRESS", email)}))

        self.assertEqual(result, [existing])

    def test_value_not_present_is_a_no_op(self) -> None:
        text = "Nessun dato sensibile qui."
        result = document_service.add_extra_value_findings(text, [], frozenset({("PERSON", "Mario Rossi")}))
        self.assertEqual(result, [])

    def test_case_sensitive_exact_match(self) -> None:
        text = "valore ABC9988 e abc9988 minuscolo"
        result = document_service.add_extra_value_findings(text, [], frozenset({("IDENTITY_DOCUMENT", "ABC9988")}))
        self.assertEqual(len(result), 1)
        self.assertEqual(text[result[0].start : result[0].end], "ABC9988")

    def test_no_extra_values_returns_findings_unchanged(self) -> None:
        findings = [Finding("EMAIL_ADDRESS", 0, 5, 0.9)]
        self.assertEqual(document_service.add_extra_value_findings("hello", findings, None), findings)
        self.assertEqual(document_service.add_extra_value_findings("hello", findings, frozenset()), findings)


class DocumentAnonymizationTest(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = PrivacyEngine()
        self.tmp = tempfile.TemporaryDirectory()
        self.base = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_anonymizes_txt_and_docx_documents(self) -> None:
        txt_path = self.base / "test.txt"
        txt_path.write_text("Il sottoscritto Mario Rossi lavora per Alfa Beta S.r.l.", encoding="utf-8")

        docx_path = self.base / "test.docx"
        doc = Document()
        doc.add_paragraph("signora Maria Bianchi email maria@example.com")
        doc.save(docx_path)

        pdf_path = self.base / "test.pdf"
        pdf = canvas.Canvas(str(pdf_path))
        pdf.drawString(72, 720, "società Rossi Consulting spa tel +39 333 123 4567")
        pdf.save()

        for path in (txt_path, docx_path):
            loaded = load_document(path)
            result = anonymize_loaded_document(loaded, self.engine)
            self.assertGreaterEqual(len(result.findings), 1)
            self.assertTrue(result.data)

        loaded_pdf = load_document(pdf_path)
        self.assertIn("Rossi Consulting", loaded_pdf.text)
        pdf_result = anonymize_loaded_document(loaded_pdf, self.engine, mode="maximum")
        self.assertEqual(pdf_result.filename, "test_anonimizzato.pdf")
        self.assertTrue(pdf_result.data.startswith(b"%PDF"))

    def test_pdf_anonymization_rasterizes_layout_and_removes_extractable_text(self) -> None:
        pdf_path = self.base / "contratto.pdf"
        pdf = canvas.Canvas(str(pdf_path), pagesize=(612, 792))
        pdf.setFont("Helvetica", 14)
        pdf.drawString(72, 720, "Il sottoscritto Mario Rossi email mario@example.com")
        pdf.drawString(72, 690, "Documento d'identita n. CA12345AA e targa AB123CD")
        pdf.save()

        result = anonymize_loaded_document(load_document(pdf_path), self.engine, mode="maximum")
        out_pdf = self.base / result.filename
        out_pdf.write_bytes(result.data)
        reader = PdfReader(out_pdf)
        extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)

        self.assertEqual(len(reader.pages), 1)
        self.assertEqual(float(reader.pages[0].mediabox.width), 612)
        self.assertEqual(float(reader.pages[0].mediabox.height), 792)
        self.assertIn("<PERSONA>", result.text)
        self.assertIn("<EMAIL>", result.text)
        self.assertIn("<DOCUMENTO_IDENTITA>", result.text)
        self.assertIn("<TARGA_VEICOLO>", result.text)
        self.assertNotIn("Mario Rossi", extracted_text)
        self.assertNotIn("mario@example.com", extracted_text)
        self.assertNotIn("CA12345AA", extracted_text)
        self.assertNotIn("AB123CD", extracted_text)

    def test_docx_anonymization_preserves_run_formatting(self) -> None:
        docx_path = self.base / "formatted.docx"
        doc = Document()
        paragraph = doc.add_paragraph("Il sottoscritto ")
        first_name = paragraph.add_run("Mario")
        first_name.bold = True
        last_name = paragraph.add_run(" Rossi")
        last_name.italic = True
        paragraph.add_run(" lavora in Via Appia.")
        doc.save(docx_path)

        result = anonymize_loaded_document(load_document(docx_path), self.engine)
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)

        output_paragraph = Document(out_docx).paragraphs[0]
        self.assertIn("M. R.", output_paragraph.text)
        self.assertIn("V. A.", output_paragraph.text)
        self.assertTrue(any(run.text == "M. R." and run.bold for run in output_paragraph.runs))

    def test_docx_maximum_protection_uses_full_placeholders(self) -> None:
        docx_path = self.base / "maximum.docx"
        doc = Document()
        doc.add_paragraph("Il sottoscritto Mario Rossi nato il 10/01/1980.")
        doc.save(docx_path)

        result = anonymize_loaded_document(load_document(docx_path), self.engine, mode="maximum")
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)

        output_text = Document(out_docx).paragraphs[0].text
        self.assertIn("<PERSONA>", output_text)
        self.assertIn("<DATA>", output_text)
        self.assertNotIn("Mario Rossi", output_text)
        self.assertNotIn("10/01/1980", output_text)

    def test_csv_documents_keep_csv_extension(self) -> None:
        csv_path = self.base / "clienti.csv"
        csv_path.write_text("nome,email\ncliente Mario Rossi,mario@example.com\n", encoding="utf-8")

        result = anonymize_loaded_document(load_document(csv_path), self.engine, mode="maximum")

        self.assertEqual(result.filename, "clienti_anonimizzato.csv")
        decoded = result.data.decode("utf-8")
        self.assertIn("<EMAIL>", decoded)
        self.assertNotIn("mario@example.com", decoded)

    def test_txt_uses_prefiltered_findings_without_reanalyzing(self) -> None:
        txt_path = self.base / "revisione.txt"
        text = "Il sottoscritto Mario Rossi email mario@example.com nato il 10/01/1980."
        txt_path.write_text(text, encoding="utf-8")

        loaded = load_document(txt_path)
        all_findings = self.engine.analyze(text, "maximum")
        only_email = [finding for finding in all_findings if finding.entity_type == "EMAIL_ADDRESS"]
        self.assertGreater(len(all_findings), len(only_email))

        result = anonymize_loaded_document(loaded, self.engine, mode="maximum", findings=only_email)
        decoded = result.data.decode("utf-8")

        self.assertEqual(result.findings, only_email)
        self.assertIn("<EMAIL>", decoded)
        self.assertIn("Mario Rossi", decoded)
        self.assertIn("10/01/1980", decoded)

    def test_docx_honors_value_exclusions_preserving_structure_and_formatting(self) -> None:
        email = "mario@example.com"
        docx_path = self.base / "esclusioni.docx"
        doc = Document()
        paragraph = doc.add_paragraph("Il sottoscritto ")
        bold_run = paragraph.add_run("Mario Rossi")
        bold_run.bold = True
        italic_run = paragraph.add_run(" con recapito ")
        italic_run.italic = True
        paragraph.add_run(email)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Recapito:"
        table.rows[0].cells[1].text = email
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = f"Il sottoscritto Mario Rossi, email {email}"
        doc.save(docx_path)
        source_doc = Document(docx_path)

        result = anonymize_loaded_document(
            load_document(docx_path),
            self.engine,
            mode="maximum",
            excluded_values=frozenset({("EMAIL_ADDRESS", email)}),
        )
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        output_doc = Document(out_docx)

        # (a) Excluded email in clear everywhere; name anonymized in body and header.
        output_paragraph = output_doc.paragraphs[0]
        self.assertEqual(output_paragraph.text, f"Il sottoscritto <PERSONA> con recapito {email}")
        self.assertEqual(output_doc.tables[0].rows[0].cells[1].text, email)
        header_text = output_doc.sections[0].header.paragraphs[0].text
        self.assertIn(email, header_text)
        self.assertIn("<PERSONA>", header_text)
        self.assertNotIn("Mario Rossi", header_text)

        # (b) Run formatting preserved on the anonymized and untouched runs.
        self.assertTrue(any(run.bold and run.text == "<PERSONA>" for run in output_paragraph.runs))
        self.assertTrue(any(run.italic and run.text == " con recapito " for run in output_paragraph.runs))

        # (c) Structure intact: paragraphs, table shape, run count.
        self.assertEqual(len(output_doc.paragraphs), len(source_doc.paragraphs))
        self.assertEqual(len(output_doc.tables[0].rows), len(source_doc.tables[0].rows))
        self.assertEqual(len(output_doc.tables[0].rows[0].cells), len(source_doc.tables[0].rows[0].cells))
        self.assertEqual(len(output_paragraph.runs), len(source_doc.paragraphs[0].runs))

        # (d) Raw XML: no included value left anywhere, excluded value still present.
        with zipfile.ZipFile(BytesIO(result.data)) as package:
            header_parts = [name for name in package.namelist() if name.startswith("word/header")]
            self.assertTrue(header_parts)
            for part_name in ["word/document.xml", *header_parts]:
                payload = package.read(part_name)
                self.assertNotIn("Mario Rossi".encode("utf-8"), payload)
                self.assertIn(email.encode("utf-8"), payload)

    def test_docx_reversible_mode_keeps_excluded_values_out_of_the_mapping(self) -> None:
        email = "mario@example.com"
        docx_path = self.base / "reversibile_esclusioni.docx"
        doc = Document()
        doc.add_paragraph(f"Il sottoscritto Mario Rossi email {email}")
        doc.save(docx_path)

        result = anonymize_loaded_document(
            load_document(docx_path),
            self.engine,
            mode="reversible",
            excluded_values=frozenset({("EMAIL_ADDRESS", email)}),
        )
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        output_text = "\n".join(paragraph.text for paragraph in Document(out_docx).paragraphs)

        self.assertIn("<PERSONA_1>", output_text)
        self.assertIn(email, output_text)
        self.assertNotIn("Mario Rossi", output_text)
        self.assertNotIn(email, [entry.value for entry in result.reversible_mapping])
        self.assertIn("Mario Rossi", [entry.value for entry in result.reversible_mapping])
        self.assertIn(email, result.text)
        self.assertNotIn("<EMAIL_1>", result.text)

    def test_pdf_native_honors_value_exclusions_on_final_file(self) -> None:
        import pypdfium2 as pdfium

        email = "mario@example.com"
        pdf_path = self.base / "esclusioni.pdf"
        pdf = canvas.Canvas(str(pdf_path), pagesize=(612, 792))
        pdf.setFont("Helvetica", 14)
        pdf.drawString(72, 720, "Il sottoscritto Mario Rossi")
        pdf.drawString(72, 650, f"Contatto: {email}")
        pdf.save()

        result = anonymize_loaded_document(
            load_document(pdf_path),
            self.engine,
            mode="maximum",
            excluded_values=frozenset({("EMAIL_ADDRESS", email)}),
        )

        # (a) Valid PDF, same page count; (c) no extractable text (rasterized by design).
        reader = PdfReader(BytesIO(result.data))
        self.assertEqual(len(reader.pages), 1)
        self.assertTrue(result.data.startswith(b"%PDF"))
        self.assertFalse("".join(page.extract_text() or "" for page in reader.pages).strip())

        # (b) Charbox coordinates from the SOURCE pdf, pixels checked on the output render.
        source_pdf = pdfium.PdfDocument(str(pdf_path))
        try:
            text_page = source_pdf[0].get_textpage()
            page_text = text_page.get_text_range()
            name_box = self._merged_charbox(text_page, page_text, "Mario Rossi")
            email_box = self._merged_charbox(text_page, page_text, email)
        finally:
            source_pdf.close()

        self.assertLess(self._region_mean_intensity(result.data, name_box, 792), 30)
        self.assertGreater(self._region_mean_intensity(result.data, email_box, 792), 120)

    def test_pdf_ocr_exclusion_with_spacing_mismatch_stays_redacted(self) -> None:
        from PIL import Image

        image = Image.new("RGB", (20, 20), "white")
        image_data = BytesIO()
        image.save(image_data, format="PNG")
        image_data.seek(0)
        pdf_path = self.base / "scansione_esclusioni.pdf"
        pdf = canvas.Canvas(str(pdf_path), pagesize=(200, 200))
        pdf.drawImage(ImageReader(image_data), 20, 20, width=160, height=160)
        pdf.save()

        ocr_text = OcrPageText(
            text="Il sottoscritto Mario Rossi",
            words=[
                OcrWord("Mario", 16, 21, 250, 300, 295, 330),
                OcrWord("Rossi", 22, 27, 300, 300, 340, 330),
            ],
        )
        with patch.object(document_service, "_tesseract_available", return_value=True):
            with patch.object(document_service, "_ocr_pdfium_page", return_value=ocr_text):
                loaded = load_document(pdf_path)
            with patch.object(document_service, "_ocr_image", return_value=ocr_text):
                # Il valore escluso ha una spaziatura diversa da quella vista dall'OCR:
                # il confronto esatto fallisce e il dato deve restare redatto (fail-closed).
                result = anonymize_loaded_document(
                    loaded,
                    self.engine,
                    mode="maximum",
                    excluded_values=frozenset({("PERSON", "Mario  Rossi")}),
                )

        self.assertLess(self._rendered_pixel(result.data, 270, 315), 10)
        self.assertLess(self._rendered_pixel(result.data, 320, 315), 10)

    @staticmethod
    def _merged_charbox(text_page, page_text: str, value: str) -> tuple[float, float, float, float]:
        start = page_text.index(value)
        boxes = [
            text_page.get_charbox(index)
            for index in range(start, start + len(value))
            if not page_text[index].isspace()
        ]
        return (
            min(box[0] for box in boxes),
            min(box[1] for box in boxes),
            max(box[2] for box in boxes),
            max(box[3] for box in boxes),
        )

    @staticmethod
    def _region_mean_intensity(pdf_data: bytes, box: tuple[float, float, float, float], page_height: float) -> float:
        import pypdfium2 as pdfium
        from PIL import ImageStat

        document = pdfium.PdfDocument(pdf_data)
        bitmap = document[0].render(scale=2)
        try:
            image = bitmap.to_pil().convert("L")
        finally:
            bitmap.close()
            document.close()
        left, bottom, right, top = box
        crop = image.crop(
            (int(left * 2), int((page_height - top) * 2), int(right * 2) + 1, int((page_height - bottom) * 2) + 1)
        )
        return ImageStat.Stat(crop).mean[0]

    @staticmethod
    def _rendered_pixel(pdf_data: bytes, x: int, y: int) -> int:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(pdf_data)
        bitmap = document[0].render(scale=2)
        try:
            pixel = bitmap.to_pil().convert("RGB").getpixel((x, y))
        finally:
            bitmap.close()
            document.close()
        return max(pixel)

    def test_pdf_native_redacts_manual_selection_added_via_extra_values(self) -> None:
        import pypdfium2 as pdfium

        manual_value = "ABC9988"
        pdf_path = self.base / "selezione_manuale.pdf"
        pdf = canvas.Canvas(str(pdf_path), pagesize=(612, 792))
        pdf.setFont("Helvetica", 14)
        pdf.drawString(72, 720, "Il sottoscritto Mario Rossi")
        pdf.drawString(72, 690, f"Prima occorrenza: {manual_value}")
        pdf.drawString(72, 660, f"Seconda occorrenza: {manual_value}")
        pdf.save()

        result = anonymize_loaded_document(
            load_document(pdf_path),
            self.engine,
            mode="maximum",
            extra_values=frozenset({("IDENTITY_DOCUMENT", manual_value)}),
        )

        reader = PdfReader(BytesIO(result.data))
        self.assertEqual(len(reader.pages), 1)
        self.assertFalse("".join(page.extract_text() or "" for page in reader.pages).strip())

        source_pdf = pdfium.PdfDocument(str(pdf_path))
        try:
            text_page = source_pdf[0].get_textpage()
            page_text = text_page.get_text_range()
            name_box = self._merged_charbox(text_page, page_text, "Mario Rossi")
            first_start = page_text.index(manual_value)
            second_start = page_text.index(manual_value, first_start + 1)
            first_box = self._merged_charbox_at(text_page, page_text, first_start, len(manual_value))
            second_box = self._merged_charbox_at(text_page, page_text, second_start, len(manual_value))
        finally:
            source_pdf.close()

        self.assertLess(self._region_mean_intensity(result.data, name_box, 792), 30)
        self.assertLess(self._region_mean_intensity(result.data, first_box, 792), 30)
        self.assertLess(self._region_mean_intensity(result.data, second_box, 792), 30)

    def test_docx_redacts_manual_selection_in_bold_run_and_table_cell(self) -> None:
        manual_value = "ABC9988"
        docx_path = self.base / "selezione_manuale.docx"
        doc = Document()
        paragraph = doc.add_paragraph("Documento numero ")
        bold_run = paragraph.add_run(manual_value)
        bold_run.bold = True
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Riferimento:"
        table.rows[0].cells[1].text = manual_value
        doc.save(docx_path)
        source_doc = Document(docx_path)

        result = anonymize_loaded_document(
            load_document(docx_path),
            self.engine,
            mode="maximum",
            extra_values=frozenset({("IDENTITY_DOCUMENT", manual_value)}),
        )
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        output_doc = Document(out_docx)

        output_paragraph = output_doc.paragraphs[0]
        self.assertNotIn(manual_value, output_paragraph.text)
        self.assertTrue(any(run.bold and manual_value not in run.text for run in output_paragraph.runs))
        self.assertNotIn(manual_value, output_doc.tables[0].rows[0].cells[1].text)

        # Struttura preservata.
        self.assertEqual(len(output_doc.paragraphs), len(source_doc.paragraphs))
        self.assertEqual(len(output_doc.tables[0].rows[0].cells), len(source_doc.tables[0].rows[0].cells))

        with zipfile.ZipFile(BytesIO(result.data)) as package:
            payload = package.read("word/document.xml")
            self.assertNotIn(manual_value.encode("utf-8"), payload)

    def test_docx_reversible_manual_selection_enters_mapping(self) -> None:
        manual_value = "ABC9988"
        docx_path = self.base / "selezione_manuale_reversibile.docx"
        doc = Document()
        doc.add_paragraph(f"Documento numero {manual_value} rilasciato oggi.")
        doc.save(docx_path)

        result = anonymize_loaded_document(
            load_document(docx_path),
            self.engine,
            mode="reversible",
            extra_values=frozenset({("IDENTITY_DOCUMENT", manual_value)}),
        )
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        output_text = "\n".join(paragraph.text for paragraph in Document(out_docx).paragraphs)

        self.assertNotIn(manual_value, output_text)
        self.assertIn("<DOCUMENTO_IDENTITA_1>", output_text)
        self.assertIn(manual_value, [entry.value for entry in result.reversible_mapping])
        self.assertEqual(restore_text(output_text, result.reversible_mapping), f"Documento numero {manual_value} rilasciato oggi.")

    def test_docx_extra_value_absent_from_text_is_a_no_op(self) -> None:
        docx_path = self.base / "selezione_manuale_assente.docx"
        doc = Document()
        doc.add_paragraph("Nessun dato sensibile in questo paragrafo.")
        doc.save(docx_path)

        result = anonymize_loaded_document(
            load_document(docx_path),
            self.engine,
            mode="maximum",
            extra_values=frozenset({("IDENTITY_DOCUMENT", "NONPRESENTE1")}),
        )

        self.assertEqual(result.findings, [])
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        self.assertEqual(
            Document(out_docx).paragraphs[0].text, "Nessun dato sensibile in questo paragrafo."
        )

    @staticmethod
    def _merged_charbox_at(text_page, page_text: str, start: int, length: int) -> tuple[float, float, float, float]:
        boxes = [
            text_page.get_charbox(index)
            for index in range(start, start + length)
            if not page_text[index].isspace()
        ]
        return (
            min(box[0] for box in boxes),
            min(box[1] for box in boxes),
            max(box[2] for box in boxes),
            max(box[3] for box in boxes),
        )

    def test_txt_reversible_mode_returns_mapping_for_local_restore(self) -> None:
        txt_path = self.base / "reversibile.txt"
        original_text = "Il sottoscritto Mario Rossi email mario@example.com nato il 10/01/1980."
        txt_path.write_text(original_text, encoding="utf-8")

        result = anonymize_loaded_document(load_document(txt_path), self.engine, mode="reversible")
        anonymized_text = result.data.decode("utf-8")

        self.assertIn("<PERSONA_1>", anonymized_text)
        self.assertIn("<EMAIL_1>", anonymized_text)
        self.assertIn("<DATA_1>", anonymized_text)
        self.assertEqual(restore_text(anonymized_text, result.reversible_mapping), original_text)
        encrypted = encrypt_mapping(result.reversible_mapping, "password locale")
        self.assertNotIn(b"Mario Rossi", encrypted)
        self.assertEqual(restore_text(anonymized_text, decrypt_mapping(encrypted, "password locale")), original_text)

    def test_txt_reversible_mode_reuses_preloaded_map_entries(self) -> None:
        txt_path = self.base / "reversibile_precaricata.txt"
        original_text = "Il sottoscritto Mario Rossi email mario@example.com conferma."
        txt_path.write_text(original_text, encoding="utf-8")
        preloaded = (ReversibleMapEntry(placeholder="<PERSONA_1>", entity_type="PERSON", value="Mario Rossi"),)

        result = anonymize_loaded_document(
            load_document(txt_path), self.engine, mode="reversible", reversible_entries=preloaded
        )
        anonymized_text = result.data.decode("utf-8")

        self.assertIn("<PERSONA_1>", anonymized_text)
        self.assertIn("<EMAIL_1>", anonymized_text)
        self.assertIn(preloaded[0], result.reversible_mapping)
        self.assertEqual(restore_text(anonymized_text, result.reversible_mapping), original_text)

    def test_docx_reversible_mode_uses_consistent_mapping(self) -> None:
        docx_path = self.base / "reversibile.docx"
        doc = Document()
        doc.add_paragraph("Il sottoscritto Mario Rossi email mario@example.com")
        doc.add_paragraph("Bonifico intestato a Mario Rossi.")
        doc.save(docx_path)

        result = anonymize_loaded_document(load_document(docx_path), self.engine, mode="reversible")
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        output_text = "\n".join(paragraph.text for paragraph in Document(out_docx).paragraphs)

        self.assertEqual(output_text.count("<PERSONA_1>"), 2)
        self.assertIn("<EMAIL_1>", output_text)
        self.assertNotIn("Mario Rossi", output_text)
        self.assertNotIn("mario@example.com", output_text)
        self.assertEqual(restore_text(output_text, result.reversible_mapping), "\n".join(Document(docx_path).paragraphs[i].text for i in range(2)))

    def test_docx_anonymizes_table_values_using_row_context(self) -> None:
        docx_path = self.base / "table-context.docx"
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        rows = [
            ("Documento d'identita", "CA12345AA"),
            ("Passaporto", "YA1234567"),
            ("Patente", "U1234567A"),
            ("Codice destinatario SDI", "ABC1234"),
            ("Codice univoco ufficio", "A1B2C3"),
            ("Targa veicolo aziendale", "AB123CD"),
            ("Tessera sanitaria", "8038 0000 0000 0000 0000"),
            ("PEC", "azienda@pec.it"),
            ("Numero pratica", "ED 9901"),
            ("Protocollo", "12345/2024"),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        doc.save(docx_path)

        result = anonymize_loaded_document(load_document(docx_path), self.engine, mode="maximum")
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)

        output_text = "\n".join(
            cell.text
            for table in Document(out_docx).tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertEqual(output_text.count("<DOCUMENTO_IDENTITA>"), 3)
        self.assertEqual(output_text.count("<CODICE_SDI>"), 2)
        self.assertIn("<TARGA_VEICOLO>", output_text)
        self.assertIn("<TESSERA_SANITARIA>", output_text)
        self.assertIn("<PEC>", output_text)
        self.assertEqual(output_text.count("<PROTOCOLLO_PRATICA>"), 2)
        self.assertNotIn("CA12345AA", output_text)
        self.assertNotIn("YA1234567", output_text)
        self.assertNotIn("U1234567A", output_text)
        self.assertNotIn("ABC1234", output_text)
        self.assertNotIn("A1B2C3", output_text)
        self.assertNotIn("AB123CD", output_text)
        self.assertNotIn("8038 0000", output_text)
        self.assertNotIn("azienda@pec.it", output_text)
        self.assertNotIn("ED 9901", output_text)
        self.assertNotIn("12345/2024", output_text)

    def test_docx_anonymization_preserves_structure_and_static_package_parts(self) -> None:
        docx_path = self.base / "layout.docx"
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        section.header.paragraphs[0].text = "Documento intestato a Mario Rossi"
        section.footer.paragraphs[0].text = "Contatti: mario@example.com"

        title = doc.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Il sottoscritto Mario Rossi")
        title_run.bold = True
        title_run.font.size = Pt(18)

        doc.add_paragraph("signora Maria Bianchi email maria@example.com", style="List Bullet")
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Documento d'identita"
        table.rows[0].cells[1].text = "CA12345AA"
        table.rows[1].cells[0].text = "Targa veicolo aziendale"
        table.rows[1].cells[1].text = "AB123CD"
        doc.add_picture(BytesIO(base64.b64decode(_ONE_PIXEL_PNG_BASE64)), width=Inches(0.2))
        doc.save(docx_path)

        before = _zip_entries(docx_path)
        result = anonymize_loaded_document(load_document(docx_path), self.engine, mode="maximum")
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)
        after = _zip_entries(out_docx)

        for entry_name in (
            "word/styles.xml",
            "word/settings.xml",
            "word/theme/theme1.xml",
            "word/_rels/document.xml.rels",
            "word/media/image1.png",
        ):
            if entry_name in before:
                self.assertEqual(before[entry_name], after[entry_name], entry_name)

        output_doc = Document(out_docx)
        output_section = output_doc.sections[0]
        self.assertEqual(output_section.orientation, WD_ORIENT.LANDSCAPE)
        self.assertEqual(output_section.top_margin, Inches(0.7))
        self.assertEqual(output_section.bottom_margin, Inches(0.8))
        self.assertEqual(output_doc.paragraphs[0].style.name, "Title")
        self.assertEqual(output_doc.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(output_doc.paragraphs[0].runs[0].bold)
        self.assertEqual(output_doc.paragraphs[1].style.name, "List Bullet")
        self.assertEqual(output_doc.tables[0].style.name, "Table Grid")
        self.assertEqual(len(output_doc.inline_shapes), 1)

        output_xml = _docx_xml_text(out_docx)
        self.assertNotIn("Mario Rossi", output_xml)
        self.assertNotIn("Maria Bianchi", output_xml)
        self.assertNotIn("mario@example.com", output_xml)
        self.assertNotIn("maria@example.com", output_xml)
        self.assertNotIn("CA12345AA", output_xml)
        self.assertNotIn("AB123CD", output_xml)
        self.assertIn("<PERSONA>", output_doc.paragraphs[0].text)
        self.assertIn("<EMAIL>", output_doc.paragraphs[1].text)
        self.assertIn("<DOCUMENTO_IDENTITA>", output_doc.tables[0].rows[0].cells[1].text)
        self.assertIn("<TARGA_VEICOLO>", output_doc.tables[0].rows[1].cells[1].text)

    def test_docx_anonymization_sanitizes_hidden_ooxml_and_metadata(self) -> None:
        docx_path = self.base / "hidden.docx"
        doc = Document()
        doc.add_paragraph("Relazione senza dati personali visibili.")
        doc.core_properties.author = "Mario Rossi"
        doc.core_properties.title = "Contratto Mario Rossi"
        doc.core_properties.comments = "Email mario@example.com"
        doc.save(docx_path)
        _add_hidden_docx_content(docx_path)

        result = anonymize_loaded_document(load_document(docx_path), self.engine)
        out_docx = self.base / result.filename
        out_docx.write_bytes(result.data)

        xml_text = _docx_xml_text(out_docx)
        output_doc = Document(out_docx)
        self.assertNotIn("Mario Rossi", xml_text)
        self.assertNotIn("Maria Bianchi", xml_text)
        self.assertNotIn("mario@example.com", xml_text)
        self.assertNotIn("06/12345678", xml_text)
        self.assertIn("M. R.", xml_text)
        self.assertIn("M. B.", xml_text)
        self.assertIn("&lt;EMAIL&gt;", xml_text)
        self.assertIn("&lt;TELEFONO&gt;", xml_text)
        self.assertNotIn("Mario Rossi", output_doc.core_properties.author or "")
        self.assertNotIn("Mario Rossi", output_doc.core_properties.title or "")
        self.assertNotIn("mario@example.com", output_doc.core_properties.comments or "")

    def test_rejects_image_only_pdf_before_anonymization(self) -> None:
        pdf_path = self.base / "scanned.pdf"
        image = ImageReader(BytesIO(base64.b64decode(_ONE_PIXEL_PNG_BASE64)))
        pdf = canvas.Canvas(str(pdf_path))
        pdf.drawImage(image, 72, 720, width=120, height=40)
        pdf.save()

        with patch("privacy_guardian.document_service._tesseract_available", return_value=False):
            with self.assertRaisesRegex(ValueError, "pagine: 1"):
                load_document(pdf_path)

    def test_ocr_reads_image_only_pdf_when_tesseract_is_available(self) -> None:
        pdf_path = self.base / "ocr-scanned.pdf"
        image = ImageReader(BytesIO(base64.b64decode(_ONE_PIXEL_PNG_BASE64)))
        pdf = canvas.Canvas(str(pdf_path))
        pdf.drawImage(image, 72, 720, width=120, height=40)
        pdf.save()
        ocr_text = OcrPageText(
            text="Il sottoscritto Mario Rossi email mario@example.com",
            words=[
                OcrWord("Mario", 16, 21, 20, 20, 80, 42),
                OcrWord("Rossi", 22, 27, 82, 20, 140, 42),
                OcrWord("mario@example.com", 34, 51, 150, 20, 300, 42),
            ],
        )

        with patch("privacy_guardian.document_service._tesseract_available", return_value=True):
            with patch("privacy_guardian.document_service._ocr_pdfium_page", return_value=ocr_text):
                loaded = load_document(pdf_path)
            with patch("privacy_guardian.document_service._ocr_image", return_value=ocr_text):
                result = anonymize_loaded_document(loaded, self.engine, mode="maximum")

        out_pdf = self.base / result.filename
        out_pdf.write_bytes(result.data)
        extracted_text = "\n".join(page.extract_text() or "" for page in PdfReader(out_pdf).pages)

        self.assertEqual(loaded.ocr_pages, (1,))
        self.assertIn("Mario Rossi", loaded.text)
        self.assertIn("<PERSONA>", result.text)
        self.assertIn("<EMAIL>", result.text)
        self.assertNotIn("Mario Rossi", extracted_text)
        self.assertNotIn("mario@example.com", extracted_text)

    def test_ocr_temp_image_is_reopenable_by_path_and_cleaned_up(self) -> None:
        from PIL import Image

        from privacy_guardian import document_service

        captured: dict[str, object] = {}

        def fake_run_tesseract(command: str, image_path: str) -> OcrPageText:
            captured["path"] = image_path
            captured["data"] = Path(image_path).read_bytes()
            return OcrPageText(text="", words=[])

        image = Image.new("RGB", (4, 4), "white")
        with patch("privacy_guardian.document_service._tesseract_command", return_value="tesseract"):
            with patch("privacy_guardian.document_service._run_tesseract_tsv", side_effect=fake_run_tesseract):
                document_service._ocr_image(image)

        self.assertTrue(bytes(captured["data"]).startswith(b"\x89PNG"))
        self.assertFalse(Path(str(captured["path"])).exists())

    def test_rejects_pdf_without_extractable_text(self) -> None:
        pdf_path = self.base / "blank.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with pdf_path.open("wb") as output:
            writer.write(output)

        with self.assertRaisesRegex(ValueError, "non contiene testo estraibile"):
            load_document(pdf_path)

    def test_pdf_reversible_mode_is_rejected_with_clear_message(self) -> None:
        pdf_path = self.base / "contratto.pdf"
        pdf = canvas.Canvas(str(pdf_path))
        pdf.drawString(72, 720, "Il sottoscritto Mario Rossi email mario@example.com")
        pdf.save()

        with self.assertRaisesRegex(ValueError, "modalità reversibile non è disponibile per i PDF"):
            anonymize_loaded_document(load_document(pdf_path), self.engine, mode="reversible")

    @unittest.skipUnless(Path("/usr/bin/textutil").exists(), "textutil is required for .doc support")
    def test_accepts_legacy_doc_and_outputs_docx(self) -> None:
        source_txt = self.base / "legacy-source.txt"
        source_txt.write_text("Il sottoscritto Mario Rossi lavora per Alfa Beta S.r.l.", encoding="utf-8")
        doc_path = self.base / "legacy.doc"
        subprocess.run(
            ["/usr/bin/textutil", "-convert", "doc", "-output", str(doc_path), str(source_txt)],
            check=True,
            capture_output=True,
        )

        result = anonymize_loaded_document(load_document(doc_path), self.engine)

        self.assertEqual(result.filename, "legacy_anonimizzato.docx")
        self.assertIn("M. R.", result.text)
        self.assertTrue(result.data)


class WebAppTest(unittest.TestCase):
    def test_health_reports_ner_active_flag(self) -> None:
        payload = asyncio.run(health_endpoint())

        self.assertIn("ner_active", payload)
        self.assertIsInstance(payload["ner_active"], bool)
        self.assertEqual(payload["ner_active"], web_app_engine.ner_active)

    def test_rejects_oversized_text_with_clear_message(self) -> None:
        payload = TextPayload(text="a" * (MAX_TEXT_LENGTH + 1), mode="standard")

        with self.assertRaises(HTTPException) as context:
            asyncio.run(anonymize_text_endpoint(payload))

        self.assertEqual(context.exception.status_code, 413)
        self.assertIn("Testo troppo lungo", context.exception.detail)

    def test_anonymizes_uploaded_document_and_returns_download_payload(self) -> None:
        upload = UploadFile(
            BytesIO(b"Il sottoscritto Mario Rossi email mario@example.com"),
            filename="contratto.txt",
        )
        payload = asyncio.run(anonymize_document(mode="maximum", file=upload))

        decoded = base64.b64decode(payload["content_base64"]).decode("utf-8")
        self.assertEqual(payload["filename"], "contratto_anonimizzato.txt")
        self.assertIn("<PERSONA>", decoded)
        self.assertIn("<EMAIL>", decoded)
        self.assertNotIn("Mario Rossi", decoded)
        self.assertIn("report", payload)
        self.assertTrue(all("label" in finding for finding in payload["findings"]))
        self.assertTrue(all("source_label" in finding for finding in payload["findings"]))

    def test_analyzes_uploaded_document_without_downloading_result(self) -> None:
        upload = UploadFile(
            BytesIO(b"Documento d'identita n. CA12345AA e targa veicolo aziendale AB123CD"),
            filename="controllo.txt",
        )
        payload = asyncio.run(analyze_document(mode="maximum", file=upload))

        self.assertEqual(payload["filename"], "controllo.txt")
        self.assertNotIn("content_base64", payload)
        self.assertEqual(payload["report"]["counts"]["IDENTITY_DOCUMENT"], 1)
        self.assertEqual(payload["report"]["counts"]["VEHICLE_PLATE"], 1)

    def test_web_findings_include_human_readable_labels(self) -> None:
        payload = asyncio.run(
            anonymize_text_endpoint(
                TextPayload(text="Carta d'identità n. CA12345AA e targa AB123CD.", mode="maximum")
            )
        )

        findings_by_type = {finding["entity_type"]: finding for finding in payload["findings"]}
        self.assertEqual(findings_by_type["IDENTITY_DOCUMENT"]["label"], "documento d'identità")
        self.assertEqual(findings_by_type["VEHICLE_PLATE"]["label"], "targa veicolo")
        self.assertEqual(findings_by_type["VEHICLE_PLATE"]["source_label"], "Regole italiane locali")

    def test_web_reversible_mode_requires_passphrase(self) -> None:
        with self.assertRaises(HTTPException) as context:
            asyncio.run(
                anonymize_text_endpoint(
                    TextPayload(text="Il sottoscritto Mario Rossi email mario@example.com", mode="reversible")
                )
            )

        self.assertEqual(context.exception.status_code, 400)
        self.assertIn("solo nell'app desktop", context.exception.detail)

    def test_rejects_oversized_uploaded_document(self) -> None:
        upload = UploadFile(BytesIO(b"123456789"), filename="troppo.txt")

        with patch("privacy_guardian.web_app.MAX_FILE_BYTES", 8):
            with self.assertRaises(HTTPException) as context:
                asyncio.run(anonymize_document(mode="standard", file=upload))

        self.assertEqual(context.exception.status_code, 413)
        self.assertIn("File troppo grande", context.exception.detail)

    def test_middleware_sets_security_headers_and_requires_content_length(self) -> None:
        from types import SimpleNamespace

        from starlette.responses import Response

        from privacy_guardian.web_app import privacy_headers

        async def call_next(request):
            return Response("ok")

        page_request = SimpleNamespace(url=SimpleNamespace(path="/"), method="GET", headers={})
        response = asyncio.run(privacy_headers(page_request, call_next))
        self.assertIn("default-src 'self'", response.headers["Content-Security-Policy"])
        self.assertEqual(response.headers["X-Frame-Options"], "DENY")
        self.assertEqual(response.headers["Cache-Control"], "no-store, max-age=0")

        chunked_request = SimpleNamespace(url=SimpleNamespace(path="/api/anonymize"), method="POST", headers={})
        response = asyncio.run(privacy_headers(chunked_request, call_next))
        self.assertEqual(response.status_code, 411)

    def test_anonymizes_uploaded_pdf_as_rasterized_redacted_document(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "contratto.pdf"
            pdf = canvas.Canvas(str(pdf_path))
            pdf.drawString(72, 720, "Il sottoscritto Mario Rossi email mario@example.com")
            pdf.save()
            upload = UploadFile(BytesIO(pdf_path.read_bytes()), filename="contratto.pdf")

            payload = asyncio.run(anonymize_document(mode="maximum", file=upload))
            decoded = base64.b64decode(payload["content_base64"])
            out_pdf = Path(tmpdir) / payload["filename"]
            out_pdf.write_bytes(decoded)
            extracted_text = "\n".join(page.extract_text() or "" for page in PdfReader(out_pdf).pages)

            self.assertEqual(payload["filename"], "contratto_anonimizzato.pdf")
            self.assertEqual(payload["media_type"], "application/pdf")
            self.assertTrue(decoded.startswith(b"%PDF"))
            self.assertEqual(payload["report"]["counts"]["PERSON"], 1)
            self.assertEqual(payload["report"]["counts"]["EMAIL_ADDRESS"], 1)
            self.assertNotIn("Mario Rossi", extracted_text)
            self.assertNotIn("mario@example.com", extracted_text)


class ActivityLogTest(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = PrivacyEngine()
        self.tmp = tempfile.TemporaryDirectory()
        self.base = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_activity_log_keeps_metadata_without_detected_values(self) -> None:
        source = self.base / "contratto.txt"
        source.write_text("Il sottoscritto Mario Rossi email mario@example.com", encoding="utf-8")
        findings = self.engine.analyze(source.read_text(encoding="utf-8"), mode="maximum")
        log_path = self.base / "activity.jsonl"

        entry = build_activity_entry(
            action="anonymization",
            source_kind="document",
            mode="maximum",
            findings=findings,
            source_path=source,
            output_data=b"<PERSONA> <EMAIL>",
            app_version="test",
        )
        record_activity(entry, log_path)
        entries = load_activity_entries(log_path)
        raw_log = log_path.read_text(encoding="utf-8")

        self.assertEqual(len(entries), 1)
        self.assertEqual(entries[0]["action_label"], "Anonimizzazione")
        self.assertEqual(entries[0]["source_extension"], ".txt")
        self.assertEqual(entries[0]["finding_counts"]["PERSON"], 1)
        self.assertEqual(entries[0]["finding_counts"]["EMAIL_ADDRESS"], 1)
        self.assertIsNotNone(entries[0]["source_sha256"])
        self.assertIsNotNone(entries[0]["output_sha256"])
        self.assertNotIn("Mario Rossi", raw_log)
        self.assertNotIn("mario@example.com", raw_log)
        self.assertNotIn("<PERSONA>", raw_log)

    def test_activity_log_exports_csv(self) -> None:
        log_path = self.base / "activity.jsonl"
        csv_path = self.base / "activity.csv"
        entry = build_activity_entry(
            action="analysis",
            source_kind="pasted_text",
            mode="standard",
            findings=[],
            app_version="test",
        )
        record_activity(entry, log_path)

        export_activity_log_csv(csv_path, log_path)
        csv_text = csv_path.read_text(encoding="utf-8")

        self.assertIn("timestamp,action_label,source_label", csv_text)
        self.assertIn("Analisi", csv_text)
        self.assertIn("Testo incollato", csv_text)


class ReversibleMapErrorTest(unittest.TestCase):
    def setUp(self) -> None:
        self.mapping = (ReversibleMapEntry(placeholder="<PERSONA_1>", entity_type="PERSON", value="Mario Rossi"),)

    def test_encrypt_mapping_rejects_empty_passphrase(self) -> None:
        with self.assertRaises(ReversibleMapError):
            encrypt_mapping(self.mapping, "   ")

    def test_decrypt_mapping_rejects_wrong_password(self) -> None:
        encrypted = encrypt_mapping(self.mapping, "password corretta")

        with self.assertRaises(ReversibleMapError):
            decrypt_mapping(encrypted, "password sbagliata")

    def test_decrypt_mapping_rejects_corrupted_or_truncated_data(self) -> None:
        encrypted = encrypt_mapping(self.mapping, "password locale")

        with self.assertRaises(ReversibleMapError):
            decrypt_mapping(encrypted[: len(encrypted) // 2], "password locale")

        with self.assertRaises(ReversibleMapError):
            decrypt_mapping(b"non e' un json valido", "password locale")

    def test_decrypt_mapping_rejects_unsupported_future_schema_version(self) -> None:
        encrypted = encrypt_mapping(self.mapping, "password locale")
        envelope = json.loads(encrypted.decode("utf-8"))
        envelope["schema_version"] = MAP_SCHEMA_VERSION + 1
        future_encrypted = json.dumps(envelope, ensure_ascii=False, indent=2).encode("utf-8")

        with self.assertRaises(ReversibleMapError):
            decrypt_mapping(future_encrypted, "password locale")

    def test_write_and_read_encrypted_mapping_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "mappa.omissis-map"
            write_encrypted_mapping(path, self.mapping, "password locale")
            restored = read_encrypted_mapping(path, "password locale")

        self.assertEqual(restored, self.mapping)

    def test_read_encrypted_mapping_rejects_wrong_password(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "mappa.omissis-map"
            write_encrypted_mapping(path, self.mapping, "password corretta")

            with self.assertRaises(ReversibleMapError):
                read_encrypted_mapping(path, "password sbagliata")


_ONE_PIXEL_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


def _add_hidden_docx_content(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        files = {name: source.read(name) for name in source.namelist()}

    content_types = files["[Content_Types].xml"].decode("utf-8")
    if "/word/comments.xml" not in content_types:
        content_types = content_types.replace(
            "</Types>",
            '<Override PartName="/word/comments.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/></Types>',
        )
    files["[Content_Types].xml"] = content_types.encode("utf-8")

    rels = files["word/_rels/document.xml.rels"].decode("utf-8")
    if "relationships/comments" not in rels:
        rels = rels.replace(
            "</Relationships>",
            '<Relationship Id="rIdHiddenComments" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
            'Target="comments.xml"/></Relationships>',
        )
    files["word/_rels/document.xml.rels"] = rels.encode("utf-8")

    document_xml = files["word/document.xml"].decode("utf-8")
    hidden_textbox = (
        '<w:p><w:r><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml">'
        "<v:textbox><w:txbxContent><w:p><w:r>"
        "<w:t>Textbox signora Maria Bianchi tel 06/12345678</w:t>"
        "</w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    document_xml = document_xml.replace("</w:body>", f"{hidden_textbox}</w:body>")
    files["word/document.xml"] = document_xml.encode("utf-8")

    files["word/comments.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:comment w:id="0" w:author="Mario Rossi" w:initials="MR">'
        "<w:p><w:r><w:t>Commento di Mario</w:t></w:r>"
        "<w:r><w:t> Rossi email mario@example.com</w:t></w:r></w:p>"
        "</w:comment></w:comments>"
    ).encode("utf-8")

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as target:
        for name, data in files.items():
            target.writestr(name, data)


def _docx_xml_text(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.endswith(".xml")
        )


def _zip_entries(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path, "r") as archive:
        return {name: archive.read(name) for name in archive.namelist()}


if __name__ == "__main__":
    unittest.main()
