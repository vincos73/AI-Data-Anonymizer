#!/usr/bin/env python3
"""Crea il kit DOCX condivisibile con i beta tester di OMISSIS."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "beta" / "Kit_beta_tester_OMISSIS.docx"
BLUE = RGBColor(7, 39, 67)
HEADING_BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(88, 101, 115)
ORANGE = "FDE9D2"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: tuple[int, ...], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, *, size: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(text))


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(text))


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, (9360,))
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"{title}. ")
    set_run(run, bold=True, color=BLUE)
    set_run(paragraph.add_run(text), color=BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=8.5, color=MUTED)


def add_form_row(table, label: str, value: str = "") -> None:
    row = table.add_row()
    set_table_geometry(table, (2500, 6860))
    row.cells[0].text = label
    row.cells[1].text = value or " "
    shade_cell(row.cells[0], LIGHT_GRAY)
    for run in row.cells[0].paragraphs[0].runs:
        set_run(run, bold=True, color=BLUE)
    for run in row.cells[1].paragraphs[0].runs:
        set_run(run)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(24, 37, 49)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, HEADING_BLUE, 18, 10),
        "Heading 2": (13, HEADING_BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Kit tester | pagina "), size=8.5, color=MUTED)
    add_page_field(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(6)
    set_run(kicker.add_run("PROGRAMMA DI TEST ESTERNO"), size=9, bold=True, color=HEADING_BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_run(title.add_run("Prova OMISSIS"), size=26, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run(subtitle.add_run("Guida operativa e scheda feedback per la closed beta"), size=14, color=MUTED)

    metadata = doc.add_table(rows=2, cols=2)
    metadata.cell(0, 0).text = "Durata prevista\n45-60 minuti"
    metadata.cell(0, 1).text = "Materiale\n3 documenti sintetici"
    metadata.cell(1, 0).text = "Dati reali\nRestano sul tuo computer"
    metadata.cell(1, 1).text = "Invio feedback\nSenza dati personali"
    set_table_geometry(metadata, (4680, 4680))
    for row in metadata.rows:
        for cell in row.cells:
            shade_cell(cell, LIGHT_BLUE)
            for index, paragraph in enumerate(cell.paragraphs):
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run(run, bold=index == 0, color=BLUE)

    add_callout(
        doc,
        "Regola principale",
        "non inviare al gruppo documenti originali, testi reali, password o File di ripristino. Se devi mostrare un problema, ricrealo con dati inventati.",
        ORANGE,
    )

    add_heading(doc, "Che cosa stiamo verificando", 1)
    doc.add_paragraph(
        "Vogliamo capire se OMISSIS riconosce i dati personali, mantiene leggibile il contenuto utile e permette di salvare una copia protetta senza assistenza. Ci interessa anche ciò che non funziona: un dato sfuggito, un oscuramento eccessivo o un passaggio poco chiaro è una segnalazione utile."
    )

    add_heading(doc, "Prova principale", 1)
    for step in (
        "Installa e avvia la versione indicata nell'invito. Annota eventuali avvisi o passaggi inattesi.",
        "Carica il primo documento sintetico assegnato.",
        "Seleziona Massima protezione e avvia l'analisi.",
        "Controlla le evidenziazioni, prestando attenzione a nomi, recapiti, indirizzi, identificativi e numeri di pratica.",
        "Leggi il risultato: le informazioni utili al significato dovrebbero restare comprensibili.",
        "Salva la copia protetta e riaprila con l'applicazione abituale.",
        "Se il documento è un PDF, prova a cercare o copiare uno dei valori originali. Non deve essere recuperabile.",
        "Ripeti con gli altri due documenti e compila una scheda per ogni problema distinto.",
    ):
        add_numbered(doc, step)

    doc.add_page_break()
    add_heading(doc, "Prove aggiuntive", 1)
    add_heading(doc, "Modalità Standard", 2)
    doc.add_paragraph(
        "Se ti è stata assegnata, ripeti un caso in modalità Standard. È normale che mantenga più contesto, per esempio iniziali e date. Segnala ciò che permette di riconoscere una persona in modo inatteso o ciò che viene oscurato senza motivo."
    )
    add_heading(doc, "Modalità Reversibile", 2)
    for step in (
        "Usa un file TXT o DOCX e scegli Reversibile.",
        "Salva il File di ripristino quando richiesto.",
        "Copia o salva il testo protetto.",
        "Incolla nel passaggio finale una breve risposta di prova che conservi gli stessi segnaposti.",
        "Ricostruisci il testo e controlla che i valori originali tornino nel posto giusto.",
    ):
        add_numbered(doc, step)

    add_callout(
        doc,
        "Prova facoltativa con un documento tuo",
        "puoi farla localmente, ma il file e il testo non devono essere inviati. Nel feedback indica soltanto il tipo di documento e la categoria del dato interessato.",
    )

    add_heading(doc, "Quando compilare una segnalazione", 1)
    for item in (
        "un dato personale resta visibile o recuperabile nel file esportato;",
        "viene oscurata una frase innocua o necessaria a capire il documento;",
        "il file salvato è illeggibile, corrotto o diverso dall'originale;",
        "non è chiaro quale azione compiere;",
        "l'app si blocca, mostra un errore o non completa l'operazione;",
        "l'installazione richiede passaggi non descritti.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Scheda feedback", 1)
    doc.add_paragraph("Compila una scheda per ogni problema distinto. Non inserire dati personali reali.")

    table = doc.add_table(rows=0, cols=2)
    add_form_row(table, "ID tester")
    add_form_row(table, "Data")
    add_form_row(table, "Versione OMISSIS")
    add_form_row(table, "Sistema operativo e versione")
    add_form_row(table, "ID del caso sintetico")
    add_form_row(table, "Formato", "☐ TXT   ☐ CSV   ☐ DOCX   ☐ PDF nativo   ☐ PDF scansionato")
    add_form_row(table, "Modalità", "☐ Standard   ☐ Massima protezione   ☐ Reversibile")
    add_form_row(table, "Attività completata", "☐ Sì   ☐ No")
    add_form_row(table, "Gravità", "☐ B0 privacy   ☐ B1 blocco d'uso   ☐ M2 malfunzionamento   ☐ M3 miglioria")

    add_heading(doc, "Descrizione del problema", 2)
    doc.add_paragraph("Che cosa è successo? Usa solo dati inventati o categorie generiche.")
    for _ in range(3):
        paragraph = doc.add_paragraph("________________________________________________________________________________")
        paragraph.paragraph_format.space_after = Pt(5)

    add_heading(doc, "Risultato atteso", 2)
    for _ in range(2):
        paragraph = doc.add_paragraph("________________________________________________________________________________")
        paragraph.paragraph_format.space_after = Pt(5)

    add_heading(doc, "Passaggi per riprodurre il problema", 2)
    for _ in range(3):
        paragraph = doc.add_paragraph("________________________________________________________________________________")
        paragraph.paragraph_format.space_after = Pt(5)

    add_callout(
        doc,
        "Allegati",
        "puoi aggiungere uno screenshot o un file soltanto dopo aver verificato che non contenga nomi, recapiti, identificativi, percorsi locali, password o dati provenienti da documenti reali.",
        ORANGE,
    )

    doc.core_properties.title = "Kit beta tester OMISSIS"
    doc.core_properties.subject = "Guida e scheda feedback per la closed beta"
    doc.core_properties.author = "OMISSIS"
    doc.core_properties.last_modified_by = "OMISSIS"
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
