#!/usr/bin/env python3
"""Genera un corpus interamente sintetico per la valutazione manuale di OMISSIS."""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "documenti_di_prova" / "corpus_multisettore"
SYNTHETIC_NOTICE = "DATI INTERAMENTE FITTIZI - DOCUMENTO DI TEST OMISSIS"


@dataclass(frozen=True)
class ExpectedValue:
    entity_type: str
    value: str


@dataclass(frozen=True)
class Case:
    case_id: str
    sector: str
    title: str
    filename: str
    text: str
    expected_remove: tuple[ExpectedValue, ...]
    must_remain: tuple[str, ...]
    tags: tuple[str, ...] = ()
    csv_rows: tuple[tuple[str, ...], ...] = ()
    csv_header: tuple[str, ...] = ()
    metadata: dict[str, str] = field(default_factory=dict)


def ev(entity_type: str, value: str) -> ExpectedValue:
    return ExpectedValue(entity_type, value)


def build_cases() -> list[Case]:
    return [
        Case(
            "SAN-01", "Sanità", "Referto di laboratorio", "SAN-01_referto_laboratorio.docx",
            """Assistita: Giulia Ferri\nCodice fiscale: FRRGLI85C45F205X\nResidenza: Via delle Magnolie 17, 20121 Milano\nEmail: giulia.ferri@example.com\nTelefono: +39 333 410 7286\nTessera sanitaria: 8038 0000 1111 2222 3333\n\nEsito: valori glicemici nei limiti; si consiglia controllo annuale.""",
            (ev("PERSON", "Giulia Ferri"), ev("CODICE_FISCALE", "FRRGLI85C45F205X"), ev("ADDRESS", "Via delle Magnolie 17"), ev("EMAIL_ADDRESS", "giulia.ferri@example.com"), ev("PHONE_NUMBER", "+39 333 410 7286"), ev("HEALTH_CARD", "8038 0000 1111 2222 3333")),
            ("valori glicemici", "controllo annuale"), ("sanitario", "docx"),
        ),
        Case(
            "SAN-02", "Sanità", "Lettera di dimissione", "SAN-02_lettera_dimissione.pdf",
            """Paziente: Luca Moretti\nNato a Bologna il 14/09/1977\nCodice fiscale: MRTLCU77P14A944Z\nDomicilio: Corso Mazzini 42, 40121 Bologna\nRecapito: 051 555 0198\n\nDiagnosi alla dimissione: polmonite risolta. Terapia domiciliare per sette giorni.""",
            (ev("PERSON", "Luca Moretti"), ev("LOCATION", "Bologna"), ev("DATE", "14/09/1977"), ev("CODICE_FISCALE", "MRTLCU77P14A944Z"), ev("ADDRESS", "Corso Mazzini 42"), ev("PHONE_NUMBER", "051 555 0198")),
            ("polmonite risolta", "Terapia domiciliare"), ("sanitario", "pdf_nativo"),
        ),
        Case(
            "SAN-03", "Sanità", "Elenco prenotazioni", "SAN-03_prenotazioni.csv",
            "Nome,Email,Telefono,Prestazione\nSara De Luca,sara.deluca@example.com,+39 347 221 9084,Visita cardiologica\nPaolo Greco,paolo.greco@example.com,+39 349 715 4432,Ecografia addominale",
            (ev("PERSON", "Sara De Luca"), ev("EMAIL_ADDRESS", "sara.deluca@example.com"), ev("PHONE_NUMBER", "+39 347 221 9084"), ev("PERSON", "Paolo Greco"), ev("EMAIL_ADDRESS", "paolo.greco@example.com"), ev("PHONE_NUMBER", "+39 349 715 4432")),
            ("Visita cardiologica", "Ecografia addominale"), ("sanitario", "csv"),
            (("Sara De Luca", "sara.deluca@example.com", "+39 347 221 9084", "Visita cardiologica"), ("Paolo Greco", "paolo.greco@example.com", "+39 349 715 4432", "Ecografia addominale")),
            ("Nome", "Email", "Telefono", "Prestazione"),
        ),
        Case(
            "SAN-04", "Psicologia", "Nota di colloquio scansionata", "SAN-04_nota_colloquio_scansione.pdf",
            """Assistita: Elena Gatti\nIndirizzo: Via del Glicine 9, 50123 Firenze\nTelefono: 055 420 1188\nEmail: elena.gatti@example.com\n\nColloquio di follow-up: la persona riferisce un miglioramento del sonno. Nessuna variazione della terapia.""",
            (ev("PERSON", "Elena Gatti"), ev("ADDRESS", "Via del Glicine 9"), ev("PHONE_NUMBER", "055 420 1188"), ev("EMAIL_ADDRESS", "elena.gatti@example.com")),
            ("miglioramento del sonno", "Nessuna variazione"), ("sanitario", "pdf_scansionato", "ocr"),
        ),
        Case(
            "LAV-01", "Lavoro e HR", "Curriculum vitae", "LAV-01_curriculum.docx",
            """Curriculum vitae di Martina Riva\nResidente in Via Manzoni 31, 10121 Torino\nEmail: martina.riva@example.com\nCellulare: +39 335 802 6194\nData di nascita: 02/12/1991\n\nCompetenze: Python, gestione progetti, comunicazione digitale. Esperienza in coordinamento di gruppi multidisciplinari.""",
            (ev("PERSON", "Martina Riva"), ev("ADDRESS", "Via Manzoni 31"), ev("EMAIL_ADDRESS", "martina.riva@example.com"), ev("PHONE_NUMBER", "+39 335 802 6194"), ev("DATE", "02/12/1991")),
            ("Python", "gestione progetti", "comunicazione digitale"), ("hr", "docx"),
        ),
        Case(
            "LAV-02", "Lavoro e HR", "Distinta pagamenti", "LAV-02_distinta_pagamenti.csv",
            "Dipendente,Codice fiscale,IBAN,Voce\nROSSI Marco,RSSMRC88A01H501T,IT60 X054 2811 1010 0000 0123 456,Rimborso trasferta",
            (ev("PERSON", "ROSSI Marco"), ev("CODICE_FISCALE", "RSSMRC88A01H501T"), ev("IBAN", "IT60 X054 2811 1010 0000 0123 456")),
            ("Rimborso trasferta",), ("hr", "bancario", "csv"),
            (("ROSSI Marco", "RSSMRC88A01H501T", "IT60 X054 2811 1010 0000 0123 456", "Rimborso trasferta"),),
            ("Dipendente", "Codice fiscale", "IBAN", "Voce"),
        ),
        Case(
            "LAV-03", "Lavoro e HR", "Nota di selezione", "LAV-03_nota_selezione.txt",
            """La candidata Elena Costa, contattabile a elena.costa@example.com e al numero 333 650 4471, ha completato il colloquio il 18/06/2026. Profilo coerente con il ruolo di responsabile operations.""",
            (ev("PERSON", "Elena Costa"), ev("EMAIL_ADDRESS", "elena.costa@example.com"), ev("PHONE_NUMBER", "333 650 4471"), ev("DATE", "18/06/2026")),
            ("responsabile operations", "Profilo coerente"), ("hr", "txt"),
        ),
        Case(
            "LAV-04", "Lavoro e HR", "Idoneità alla mansione", "LAV-04_idoneita_mansione.docx",
            """Dipendente: Andrea Villa\nCodice fiscale: VLLNDR90E15L219Q\nMatricola interna: DIP-00427\nSede di lavoro: Piazza Statuto 6, Torino\n\nGiudizio: idoneo alla mansione con prescrizione di pausa visiva periodica.""",
            (ev("PERSON", "Andrea Villa"), ev("CODICE_FISCALE", "VLLNDR90E15L219Q"), ev("ADDRESS", "Piazza Statuto 6")),
            ("DIP-00427", "idoneo alla mansione", "pausa visiva"), ("hr", "sanitario", "docx"),
        ),
        Case(
            "IST-01", "Scuola", "Comunicazione alla famiglia", "IST-01_comunicazione_famiglia.pdf",
            """Alunno: Davide Conti\nNato a Bari il 09/03/2012\nResidente in Via Dante 28, 70121 Bari\nGenitore: Chiara Conti\nEmail: chiara.conti@example.com\nTelefono: +39 340 919 2375\n\nOggetto: partecipazione al laboratorio di robotica educativa.""",
            (ev("PERSON", "Davide Conti"), ev("LOCATION", "Bari"), ev("DATE", "09/03/2012"), ev("ADDRESS", "Via Dante 28"), ev("PERSON", "Chiara Conti"), ev("EMAIL_ADDRESS", "chiara.conti@example.com"), ev("PHONE_NUMBER", "+39 340 919 2375")),
            ("laboratorio di robotica educativa",), ("scuola", "pdf_nativo"),
        ),
        Case(
            "IST-02", "Università", "Domanda di borsa di studio", "IST-02_borsa_studio.docx",
            """Richiedente: Francesca Longo\nCodice fiscale: LNGFNC01D62F839R\nEmail istituzionale: francesca.longo@example.com\nIBAN: IT60 X054 2811 1010 0000 0123 456\nIscritta all'Università degli Studi di Napoli Federico II, corso di Economia.\n\nLa domanda riguarda il contributo per mobilità internazionale.""",
            (ev("PERSON", "Francesca Longo"), ev("CODICE_FISCALE", "LNGFNC01D62F839R"), ev("EMAIL_ADDRESS", "francesca.longo@example.com"), ev("IBAN", "IT60 X054 2811 1010 0000 0123 456"), ev("ORGANIZATION", "Università degli Studi di Napoli Federico II")),
            ("corso di Economia", "mobilità internazionale"), ("universita", "docx"),
        ),
        Case(
            "FIN-01", "Amministrazione e finanza", "Fattura di consulenza", "FIN-01_fattura.pdf",
            """Fornitore: Aurora Digitale S.r.l.\nSede: Viale Europa 45, 20122 Milano\nPartita IVA: 12345678903\nPEC: auroradigitale@pec.example.com\nCodice destinatario SDI: ABC1234\nIBAN: IT60 X054 2811 1010 0000 0123 456\n\nDescrizione: consulenza organizzativa - giugno 2026. Imponibile euro 2.400,00.""",
            (ev("ORGANIZATION", "Aurora Digitale S.r.l."), ev("ADDRESS", "Viale Europa 45"), ev("PARTITA_IVA", "12345678903"), ev("PEC_ADDRESS", "auroradigitale@pec.example.com"), ev("SDI_CODE", "ABC1234"), ev("IBAN", "IT60 X054 2811 1010 0000 0123 456")),
            ("consulenza organizzativa", "Imponibile euro 2.400,00"), ("contabilita", "pdf_nativo"),
        ),
        Case(
            "FIN-02", "Banca", "Disposizione di bonifico", "FIN-02_bonifico.txt",
            """Ordinante: Matteo Serra\nCodice fiscale: SRRMTT79H04I452V\nBeneficiario: Studio Quercia S.a.s.\nIBAN beneficiario: IT60 X054 2811 1010 0000 0123 456\nCarta usata per la verifica: 4111 1111 1111 1111\nCausale: anticipo progetto editoriale.""",
            (ev("PERSON", "Matteo Serra"), ev("CODICE_FISCALE", "SRRMTT79H04I452V"), ev("ORGANIZATION", "Studio Quercia S.a.s."), ev("IBAN", "IT60 X054 2811 1010 0000 0123 456"), ev("CREDIT_CARD", "4111 1111 1111 1111")),
            ("anticipo progetto editoriale",), ("bancario", "txt"),
        ),
        Case(
            "FIN-03", "Assicurazioni", "Denuncia di sinistro", "FIN-03_denuncia_sinistro.docx",
            """Assicurato: Simone Barbieri\nResidente in Via Po 12, 00198 Roma\nPatente n. U1234567A\nTarga veicolo: AB123CD\nNumero pratica: SIN 7824/2026\nTelefono: 06 702 1189\n\nDinamica: urto a bassa velocità in area di parcheggio; nessun ferito.""",
            (ev("PERSON", "Simone Barbieri"), ev("ADDRESS", "Via Po 12"), ev("IDENTITY_DOCUMENT", "U1234567A"), ev("VEHICLE_PLATE", "AB123CD"), ev("PROTOCOL_CASE_NUMBER", "SIN 7824/2026"), ev("PHONE_NUMBER", "06 702 1189")),
            ("urto a bassa velocità", "nessun ferito"), ("assicurazioni", "docx"),
        ),
        Case(
            "IMM-01", "Condominio", "Verbale di assemblea", "IMM-01_verbale_condominio.pdf",
            """Condominio di Via Verdi 18, Genova\nPresiede l'amministratrice Laura Neri.\nProprietario dell'interno 7: Roberto Sala, email roberto.sala@example.com.\n\nL'assemblea approva il preventivo per la manutenzione dell'ascensore e rinvia la tinteggiatura del vano scale.""",
            (ev("ADDRESS", "Via Verdi 18"), ev("PERSON", "Laura Neri"), ev("PERSON", "Roberto Sala"), ev("EMAIL_ADDRESS", "roberto.sala@example.com")),
            ("manutenzione dell'ascensore", "vano scale"), ("condominio", "pdf_nativo"),
        ),
        Case(
            "IMM-02", "Immobiliare", "Scheda catastale", "IMM-02_scheda_catastale.docx",
            """Intestataria: Monica Pellegrini\nCodice fiscale: PLLMNC82R55G273W\nImmobile sito in Via dei Pini 6, Palermo\nComune catastale: Palermo\nFoglio 24, particella 318, subalterno 7, categoria catastale A/3.\n\nNota: appartamento ad uso abitativo con pertinenza.""",
            (ev("PERSON", "Monica Pellegrini"), ev("CODICE_FISCALE", "PLLMNC82R55G273W"), ev("ADDRESS", "Via dei Pini 6"), ev("CATASTO", "24"), ev("CATASTO", "318"), ev("CATASTO", "7"), ev("CATASTO", "A/3")),
            ("appartamento ad uso abitativo", "pertinenza"), ("immobiliare", "catasto", "docx"),
        ),
        Case(
            "SRV-01", "Commercio elettronico", "Conferma ordine", "SRV-01_conferma_ordine.txt",
            """Cliente: Alessio Rinaldi\nEmail: alessio.rinaldi@example.com\nTelefono: +39 338 770 2541\nConsegna: Via Roma 77, 35121 Padova\nOrdine: ORD-2026-4451\n\nProdotto: cuffie wireless; quantità 1; consegna standard.""",
            (ev("PERSON", "Alessio Rinaldi"), ev("EMAIL_ADDRESS", "alessio.rinaldi@example.com"), ev("PHONE_NUMBER", "+39 338 770 2541"), ev("ADDRESS", "Via Roma 77")),
            ("ORD-2026-4451", "cuffie wireless", "consegna standard"), ("ecommerce", "txt"),
        ),
        Case(
            "SRV-02", "Assistenza clienti", "Esportazione ticket", "SRV-02_ticket_assistenza.csv",
            "Cliente,Email,Telefono,Ticket,Problema\nBeatrice Romano,beatrice.romano@example.com,320 441 9087,TCK-21904,Accesso bloccato dopo aggiornamento",
            (ev("PERSON", "Beatrice Romano"), ev("EMAIL_ADDRESS", "beatrice.romano@example.com"), ev("PHONE_NUMBER", "320 441 9087")),
            ("TCK-21904", "Accesso bloccato dopo aggiornamento"), ("customer_care", "csv"),
            (("Beatrice Romano", "beatrice.romano@example.com", "320 441 9087", "TCK-21904", "Accesso bloccato dopo aggiornamento"),),
            ("Cliente", "Email", "Telefono", "Ticket", "Problema"),
        ),
        Case(
            "SRV-03", "Energia", "Riepilogo fornitura", "SRV-03_bolletta_energia.pdf",
            """Intestatario: Stefano Amato\nCodice fiscale: MTASFN76B11C351K\nPunto di fornitura: Via Etnea 104, 95131 Catania\nEmail: stefano.amato@example.com\nTelefono: 095 601 8821\n\nConsumo del periodo: 284 kWh. Offerta: Energia Casa Verde.""",
            (ev("PERSON", "Stefano Amato"), ev("CODICE_FISCALE", "MTASFN76B11C351K"), ev("ADDRESS", "Via Etnea 104"), ev("EMAIL_ADDRESS", "stefano.amato@example.com"), ev("PHONE_NUMBER", "095 601 8821")),
            ("284 kWh", "Energia Casa Verde"), ("utility", "pdf_nativo"),
        ),
        Case(
            "SRV-04", "Telecomunicazioni", "Reclamo connessione", "SRV-04_reclamo_telefonia.txt",
            """Il cliente Giorgio Piras, residente in Via Sardegna 22, Cagliari, segnala disconnessioni serali. Contatto: giorgio.piras@example.com, 070 440 1882. Numero linea: LINEA-77821. Richiede verifica tecnica senza modifica del piano tariffario.""",
            (ev("PERSON", "Giorgio Piras"), ev("ADDRESS", "Via Sardegna 22"), ev("EMAIL_ADDRESS", "giorgio.piras@example.com"), ev("PHONE_NUMBER", "070 440 1882")),
            ("LINEA-77821", "disconnessioni serali", "piano tariffario"), ("telecom", "txt"),
        ),
        Case(
            "TUR-01", "Turismo", "Conferma prenotazione alberghiera", "TUR-01_prenotazione_hotel.docx",
            """Ospite: Valentina Farina\nPassaporto n. YA1234567\nEmail: valentina.farina@example.com\nTelefono: +39 346 512 8804\nArrivo: 21/09/2026\nPartenza: 24/09/2026\n\nCamera doppia uso singola, colazione inclusa.""",
            (ev("PERSON", "Valentina Farina"), ev("IDENTITY_DOCUMENT", "YA1234567"), ev("EMAIL_ADDRESS", "valentina.farina@example.com"), ev("PHONE_NUMBER", "+39 346 512 8804"), ev("DATE", "21/09/2026"), ev("DATE", "24/09/2026")),
            ("Camera doppia uso singola", "colazione inclusa"), ("turismo", "docx"),
        ),
        Case(
            "TUR-02", "Viaggi", "Richiesta rimborso scansionata", "TUR-02_rimborso_viaggio_scansione.pdf",
            """Richiedente: Fabio Caruso\nCodice fiscale: CRSFBA84L22F205N\nEmail: fabio.caruso@example.com\nIBAN: IT60 X054 2811 1010 0000 0123 456\nPratica n. RIM 9138/2026\n\nMotivo: cancellazione del collegamento ferroviario del 3 luglio 2026.""",
            (ev("PERSON", "Fabio Caruso"), ev("CODICE_FISCALE", "CRSFBA84L22F205N"), ev("EMAIL_ADDRESS", "fabio.caruso@example.com"), ev("IBAN", "IT60 X054 2811 1010 0000 0123 456"), ev("PROTOCOL_CASE_NUMBER", "RIM 9138/2026"), ev("DATE", "3 luglio 2026")),
            ("cancellazione del collegamento ferroviario",), ("viaggi", "pdf_scansionato", "ocr"),
        ),
        Case(
            "PA-01", "Pubblica amministrazione", "Comunicazione protocollata", "PA-01_comunicazione_protocollo.docx",
            """Protocollo n. 18472/2026\nDestinataria: Silvia Fontana\nCodice fiscale: FNTSLV73A41L219M\nResidenza: Corso Francia 15, Torino\nPEC: silvia.fontana@pec.example.com\n\nIl Comune di Torino comunica l'avvio del procedimento relativo al contributo per l'efficienza energetica.""",
            (ev("PROTOCOL_CASE_NUMBER", "18472/2026"), ev("PERSON", "Silvia Fontana"), ev("CODICE_FISCALE", "FNTSLV73A41L219M"), ev("ADDRESS", "Corso Francia 15"), ev("PEC_ADDRESS", "silvia.fontana@pec.example.com"), ev("TERRITORIAL_BODY", "Comune di Torino")),
            ("avvio del procedimento", "efficienza energetica"), ("pa", "docx"),
        ),
        Case(
            "SOC-01", "Servizi sociali", "Relazione di presa in carico", "SOC-01_relazione_sociale.pdf",
            """Beneficiaria: Nadia Esposito\nNata a Salerno il 06/08/1988\nDomicilio: Via Mercanti 33, 84121 Salerno\nTelefono: +39 339 234 6008\nEmail: nadia.esposito@example.com\n\nIntervento proposto: orientamento al lavoro e sostegno temporaneo per la mobilità.""",
            (ev("PERSON", "Nadia Esposito"), ev("LOCATION", "Salerno"), ev("DATE", "06/08/1988"), ev("ADDRESS", "Via Mercanti 33"), ev("PHONE_NUMBER", "+39 339 234 6008"), ev("EMAIL_ADDRESS", "nadia.esposito@example.com")),
            ("orientamento al lavoro", "sostegno temporaneo"), ("sociale", "pdf_nativo"),
        ),
        Case(
            "ETS-01", "Terzo settore", "Registro donazioni", "ETS-01_registro_donazioni.csv",
            "Donatore,Email,Importo,Progetto\nIlaria Bassi,ilaria.bassi@example.com,150,Emporio solidale\nClaudio Testa,claudio.testa@example.com,75,Doposcuola di quartiere",
            (ev("PERSON", "Ilaria Bassi"), ev("EMAIL_ADDRESS", "ilaria.bassi@example.com"), ev("PERSON", "Claudio Testa"), ev("EMAIL_ADDRESS", "claudio.testa@example.com")),
            ("Emporio solidale", "Doposcuola di quartiere"), ("terzo_settore", "csv"),
            (("Ilaria Bassi", "ilaria.bassi@example.com", "150", "Emporio solidale"), ("Claudio Testa", "claudio.testa@example.com", "75", "Doposcuola di quartiere")),
            ("Donatore", "Email", "Importo", "Progetto"),
        ),
        Case(
            "SPT-01", "Sport", "Modulo di iscrizione", "SPT-01_iscrizione_sportiva.txt",
            """Atleta: Riccardo Bernardi\nNato a Verona il 17/04/2006\nCodice fiscale: BRNRCR06D17L781J\nTelefono del genitore: 345 610 9934\nEmail: famiglia.bernardi@example.com\n\nDisciplina: pallavolo; categoria Under 21; certificato medico acquisito.""",
            (ev("PERSON", "Riccardo Bernardi"), ev("LOCATION", "Verona"), ev("DATE", "17/04/2006"), ev("CODICE_FISCALE", "BRNRCR06D17L781J"), ev("PHONE_NUMBER", "345 610 9934"), ev("EMAIL_ADDRESS", "famiglia.bernardi@example.com")),
            ("pallavolo", "Under 21", "certificato medico acquisito"), ("sport", "txt"),
        ),
        Case(
            "LOG-01", "Logistica", "Documento di consegna", "LOG-01_documento_consegna.pdf",
            """Destinataria: Federica Marchetti\nConsegna presso Via Adriatica 56, 60121 Ancona\nTelefono: 071 208 4419\nEmail: federica.marchetti@example.com\nTarga mezzo: CD456EF\n\nMerce: 4 colli non deperibili. Finestra di consegna: 09:00-12:00.""",
            (ev("PERSON", "Federica Marchetti"), ev("ADDRESS", "Via Adriatica 56"), ev("PHONE_NUMBER", "071 208 4419"), ev("EMAIL_ADDRESS", "federica.marchetti@example.com"), ev("VEHICLE_PLATE", "CD456EF")),
            ("4 colli non deperibili", "09:00-12:00"), ("logistica", "pdf_nativo"),
        ),
        Case(
            "TEC-01", "Tecnologia", "Rapporto di incidente", "TEC-01_incidente_sicurezza.docx",
            """Segnalante: Pietro Fabbri\nEmail: pietro.fabbri@example.com\nTelefono: +39 334 713 0882\nSocietà: Rete Sicura S.p.A.\n\nEvento: accessi ripetuti non riusciti al portale interno. Identificativo tecnico HOST-LAB-27. Nessuna esfiltrazione rilevata.""",
            (ev("PERSON", "Pietro Fabbri"), ev("EMAIL_ADDRESS", "pietro.fabbri@example.com"), ev("PHONE_NUMBER", "+39 334 713 0882"), ev("ORGANIZATION", "Rete Sicura S.p.A.")),
            ("HOST-LAB-27", "Nessuna esfiltrazione rilevata"), ("tecnologia", "docx"),
        ),
        Case(
            "MKT-01", "Marketing e CRM", "Esportazione contatti evento", "MKT-01_contatti_evento.csv",
            "Nome,Azienda,Email,Telefono,Interesse\nCamilla Guerra,Nord Est Media S.r.l.,camilla.guerra@example.com,+39 348 120 7745,Automazione marketing",
            (ev("PERSON", "Camilla Guerra"), ev("ORGANIZATION", "Nord Est Media S.r.l."), ev("EMAIL_ADDRESS", "camilla.guerra@example.com"), ev("PHONE_NUMBER", "+39 348 120 7745")),
            ("Automazione marketing",), ("marketing", "crm", "csv"),
            (("Camilla Guerra", "Nord Est Media S.r.l.", "camilla.guerra@example.com", "+39 348 120 7745", "Automazione marketing"),),
            ("Nome", "Azienda", "Email", "Telefono", "Interesse"),
        ),
        Case(
            "PRO-01", "Professionisti", "Lettera di incarico", "PRO-01_incarico_freelance.txt",
            """Committente: Blu Progetti S.r.l., con sede in Piazza Dante 11, Napoli.\nProfessionista: Anna Vitale, codice fiscale VTLNNA86T41F839D, PEC anna.vitale@pec.example.com.\nCompenso concordato: euro 3.200 oltre accessori. L'incarico riguarda la revisione del piano editoriale.""",
            (ev("ORGANIZATION", "Blu Progetti S.r.l."), ev("ADDRESS", "Piazza Dante 11"), ev("PERSON", "Anna Vitale"), ev("CODICE_FISCALE", "VTLNNA86T41F839D"), ev("PEC_ADDRESS", "anna.vitale@pec.example.com")),
            ("euro 3.200 oltre accessori", "revisione del piano editoriale"), ("professionisti", "txt"),
        ),
        Case(
            "EVT-01", "Eventi", "Lista accrediti", "EVT-01_lista_accrediti.pdf",
            """Evento: Forum Innovazione Locale\nPartecipante: Marco De Angelis\nEmail: marco.deangelis@example.com\nTelefono: 333 908 1127\nPartecipante: Teresa Rizzi\nEmail: teresa.rizzi@example.com\n\nSessione scelta: Laboratorio servizi digitali.""",
            (ev("PERSON", "Marco De Angelis"), ev("EMAIL_ADDRESS", "marco.deangelis@example.com"), ev("PHONE_NUMBER", "333 908 1127"), ev("PERSON", "Teresa Rizzi"), ev("EMAIL_ADDRESS", "teresa.rizzi@example.com")),
            ("Forum Innovazione Locale", "Laboratorio servizi digitali"), ("eventi", "pdf_nativo"),
        ),
        Case(
            "LEG-01", "Legale", "Promemoria pratica", "LEG-01_promemoria_pratica.docx",
            """Cliente: Enrico Bellini\nCodice fiscale: BLLNRC71M09D612Y\nDomicilio: Via San Marco 14, 30124 Venezia\nFascicolo RG 4567/2026\nPEC: enrico.bellini@pec.example.com\n\nOggetto: verifica della documentazione contrattuale e predisposizione della memoria riepilogativa.""",
            (ev("PERSON", "Enrico Bellini"), ev("CODICE_FISCALE", "BLLNRC71M09D612Y"), ev("ADDRESS", "Via San Marco 14"), ev("PROTOCOL_CASE_NUMBER", "RG 4567/2026"), ev("PEC_ADDRESS", "enrico.bellini@pec.example.com")),
            ("documentazione contrattuale", "memoria riepilogativa"), ("legale", "docx"),
        ),
        Case(
            "TRP-01", "Trasporti", "Richiesta duplicato abbonamento", "TRP-01_duplicato_abbonamento.txt",
            """Richiedente: Ornella Fiore\nCodice fiscale: FRORLL68E50F839S\nEmail: ornella.fiore@example.com\nTelefono: 081 710 4406\nDocumento d'identità n. CA12345AA\n\nMotivo: tessera di viaggio deteriorata; si richiede duplicato senza variazione della tratta.""",
            (ev("PERSON", "Ornella Fiore"), ev("CODICE_FISCALE", "FRORLL68E50F839S"), ev("EMAIL_ADDRESS", "ornella.fiore@example.com"), ev("PHONE_NUMBER", "081 710 4406"), ev("IDENTITY_DOCUMENT", "CA12345AA")),
            ("tessera di viaggio deteriorata", "senza variazione della tratta"), ("trasporti", "txt"),
        ),
        Case(
            "FOR-01", "Fornitori", "Scheda di onboarding", "FOR-01_onboarding_fornitore.docx",
            """Referente: Massimo Orlando\nSocietà: Officina del Sud S.n.c.\nPartita IVA: 12345678903\nSede: Via delle Industrie 5, 70132 Bari\nPEC: officinadelsud@pec.example.com\nCodice univoco ufficio: A1B2C3\n\nCategoria: manutenzione impianti; pagamento previsto a 30 giorni.""",
            (ev("PERSON", "Massimo Orlando"), ev("ORGANIZATION", "Officina del Sud S.n.c."), ev("PARTITA_IVA", "12345678903"), ev("ADDRESS", "Via delle Industrie 5"), ev("PEC_ADDRESS", "officinadelsud@pec.example.com"), ev("SDI_CODE", "A1B2C3")),
            ("manutenzione impianti", "pagamento previsto a 30 giorni"), ("fornitori", "docx"),
        ),
        Case(
            "RET-01", "Commercio", "Elenco tessere fedeltà", "RET-01_tessere_fedelta.csv",
            "Cliente,Email,Telefono,Codice programma,Preferenza\nMarta Guidi,marta.guidi@example.com,329 703 1822,LOY-41028,Prodotti senza glutine",
            (ev("PERSON", "Marta Guidi"), ev("EMAIL_ADDRESS", "marta.guidi@example.com"), ev("PHONE_NUMBER", "329 703 1822")),
            ("LOY-41028", "Prodotti senza glutine"), ("retail", "csv"),
            (("Marta Guidi", "marta.guidi@example.com", "329 703 1822", "LOY-41028", "Prodotti senza glutine"),),
            ("Cliente", "Email", "Telefono", "Codice programma", "Preferenza"),
        ),
        Case(
            "MED-01", "Media", "Liberatoria intervista", "MED-01_liberatoria_intervista.pdf",
            """Intervistata: Rosa Coppola\nNata a Matera il 23/11/1983\nResidente in Via Lucana 63, 75100 Matera\nEmail: rosa.coppola@example.com\nTelefono: +39 331 667 2401\n\nLa partecipante autorizza l'uso dell'intervista nel podcast dedicato all'innovazione sociale.""",
            (ev("PERSON", "Rosa Coppola"), ev("LOCATION", "Matera"), ev("DATE", "23/11/1983"), ev("ADDRESS", "Via Lucana 63"), ev("EMAIL_ADDRESS", "rosa.coppola@example.com"), ev("PHONE_NUMBER", "+39 331 667 2401")),
            ("podcast dedicato all'innovazione sociale",), ("media", "pdf_nativo"),
        ),
    ]


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def create_docx(case: Case, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = f"{case.sector} | Caso {case.case_id}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 100, 112)

    notice = doc.add_paragraph()
    notice.paragraph_format.space_after = Pt(5)
    run = notice.add_run(SYNTHETIC_NOTICE)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(160, 78, 0)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(case.title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(7, 39, 67)

    blocks = case.text.split("\n\n")
    for block in blocks:
        lines = block.splitlines()
        if len(lines) > 1 and all(":" in line for line in lines):
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            for line in lines:
                label, value = line.split(":", 1)
                row = table.add_row()
                _set_cell_width(row.cells[0], 2400)
                _set_cell_width(row.cells[1], 6960)
                row.cells[0].text = label.strip()
                row.cells[1].text = value.strip()
                row.cells[0].paragraphs[0].runs[0].bold = True
                for cell in row.cells:
                    cell.vertical_alignment = 1
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(3)
                        for cell_run in para.runs:
                            cell_run.font.name = "Arial"
                            cell_run.font.size = Pt(10)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            paragraph = doc.add_paragraph(block.replace("\n", "\n"))
            paragraph.paragraph_format.space_after = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.text = "Uso esclusivo per test. Nessun dato appartiene a persone reali."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 108, 118)

    doc.core_properties.title = f"Caso sintetico {case.case_id}"
    doc.core_properties.subject = "Corpus multi-settore OMISSIS"
    doc.core_properties.author = "OMISSIS Test Lab"
    doc.core_properties.last_modified_by = "OMISSIS Test Lab"
    doc.save(path)


def create_pdf(case: Case, path: Path) -> None:
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "BodyOMISSIS",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=HexColor("#172A3A"),
        alignment=TA_LEFT,
        spaceAfter=9,
    )
    title_style = ParagraphStyle(
        "TitleOMISSIS",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=HexColor("#072743"),
        alignment=TA_LEFT,
        spaceAfter=14,
    )
    notice_style = ParagraphStyle(
        "NoticeOMISSIS",
        parent=body_style,
        fontName="Helvetica-Bold",
        fontSize=8.5,
        textColor=HexColor("#A04E00"),
        spaceAfter=7,
    )
    doc = SimpleDocTemplate(
        str(path), pagesize=A4, rightMargin=22 * mm, leftMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=f"Caso sintetico {case.case_id}", author="OMISSIS Test Lab",
    )
    story = [
        Paragraph(SYNTHETIC_NOTICE, notice_style),
        Paragraph(case.title, title_style),
        Paragraph(f"Settore: {case.sector} | Caso: {case.case_id}", body_style),
        Spacer(1, 5 * mm),
    ]
    for block in case.text.split("\n\n"):
        safe = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe.replace("\n", "<br/>"), body_style))
    story.append(Spacer(1, 8 * mm))
    story.append(Paragraph("Uso esclusivo per test. Nessun dato appartiene a persone reali.", notice_style))
    doc.build(story)


def _load_scan_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/System/Library/Fonts/Supplemental/Helvetica.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _wrap_line(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def create_scanned_pdf(case: Case, path: Path) -> None:
    width, height = 1654, 2339
    image = Image.new("RGB", (width, height), "#FDFCF8")
    draw = ImageDraw.Draw(image)
    font = _load_scan_font(34)
    title_font = _load_scan_font(48)
    small_font = _load_scan_font(27)
    x, y = 140, 120
    draw.text((x, y), SYNTHETIC_NOTICE, fill="#8A4A12", font=small_font)
    y += 80
    draw.text((x, y), case.title, fill="#102A43", font=title_font)
    y += 100
    draw.text((x, y), f"Settore: {case.sector} | Caso: {case.case_id}", fill="#52606D", font=small_font)
    y += 90
    for raw_line in case.text.splitlines():
        for line in _wrap_line(draw, raw_line, font, width - 280):
            draw.text((x, y), line, fill="#111111", font=font)
            y += 53
        if not raw_line:
            y += 22
    draw.text((x, height - 135), "Uso esclusivo per test - nessun dato reale", fill="#616E7C", font=small_font)
    image.save(path, "PDF", resolution=200.0)


def write_case(case: Case) -> None:
    path = OUTPUT_DIR / case.filename
    suffix = path.suffix.lower()
    if suffix == ".docx":
        create_docx(case, path)
    elif suffix == ".pdf" and "pdf_scansionato" in case.tags:
        create_scanned_pdf(case, path)
    elif suffix == ".pdf":
        create_pdf(case, path)
    elif suffix == ".csv":
        with path.open("w", encoding="utf-8", newline="") as stream:
            writer = csv.writer(stream, lineterminator="\n")
            writer.writerow(case.csv_header)
            writer.writerows(case.csv_rows)
    else:
        path.write_text(f"{SYNTHETIC_NOTICE}\n\n{case.text}\n", encoding="utf-8")


def write_manifest(cases: Iterable[Case]) -> None:
    payload = {
        "schema_version": 1,
        "notice": "Tutti i nomi, recapiti, identificativi e contenuti sono inventati per il test di OMISSIS.",
        "recommended_mode": "maximum",
        "cases": [
            {
                "id": case.case_id,
                "sector": case.sector,
                "title": case.title,
                "file": case.filename,
                "tags": list(case.tags),
                "expected_remove": [
                    {"entity_type": item.entity_type, "value": item.value}
                    for item in case.expected_remove
                ],
                "must_remain": (
                    ([] if case.filename.lower().endswith(".csv") else [SYNTHETIC_NOTICE])
                    + list(case.must_remain)
                ),
            }
            for case in cases
        ],
    }
    (OUTPUT_DIR / "manifest.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def write_readme(cases: list[Case]) -> None:
    formats: dict[str, int] = {}
    sectors: dict[str, int] = {}
    for case in cases:
        format_name = "PDF scansionato" if "pdf_scansionato" in case.tags else Path(case.filename).suffix.lower().lstrip(".").upper()
        formats[format_name] = formats.get(format_name, 0) + 1
        sectors[case.sector] = sectors.get(case.sector, 0) + 1
    lines = [
        "# Corpus sintetico multi-settore OMISSIS",
        "",
        f"Il corpus contiene **{len(cases)} documenti interamente fittizi**. Non contiene dati personali reali e può essere usato in demo, test automatici e sessioni con beta tester.",
        "",
        "## Come usarlo",
        "",
        "1. Aprire un documento con OMISSIS.",
        "2. Eseguire prima il test in modalità **Massima protezione**.",
        "3. Confrontare il risultato con `manifest.json`: `expected_remove` deve sparire, `must_remain` deve restare leggibile.",
        "4. Per i PDF, salvare la copia protetta e verificare che i valori originali non siano selezionabili o estraibili.",
        "5. Eseguire `PYTHONPATH=src python3 scripts/evaluate_multisector_corpus.py` per produrre una baseline ripetibile.",
        "",
        "## Copertura dei formati",
        "",
    ]
    lines.extend(f"- {name}: {count}" for name, count in sorted(formats.items()))
    lines.extend(["", "## Copertura dei settori", ""])
    lines.extend(f"- {name}: {count}" for name, count in sorted(sectors.items()))
    lines.extend([
        "",
        "## Regole di sicurezza",
        "",
        "- Non sostituire i dati fittizi con dati reali nel repository.",
        "- I tester esterni possono usare propri documenti localmente, ma non devono inviare gli originali al team.",
        "- Nei report descrivere il tipo di dato sfuggito; allegare soltanto estratti già anonimizzati o ricreati con valori inventati.",
        "- Un risultato senza segnalazioni non equivale a garanzia assoluta: resta obbligatorio il controllo umano.",
        "",
    ])
    (OUTPUT_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    cases = build_cases()
    if len(cases) != 35:
        raise RuntimeError(f"Il corpus deve contenere 35 casi, trovati: {len(cases)}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for case in cases:
        write_case(case)
    write_manifest(cases)
    write_readme(cases)
    print(f"Creati {len(cases)} casi in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
