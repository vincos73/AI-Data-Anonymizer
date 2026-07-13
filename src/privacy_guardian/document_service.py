from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import Iterable
import csv
import os
import shutil
import subprocess
import sys
import tempfile
import xml.etree.ElementTree as ET
import zipfile

from privacy_guardian.models import AnonymizationMode, Finding
from privacy_guardian.privacy_engine import PrivacyEngine
from privacy_guardian.reversible import ReversibleAnonymizer, ReversibleMapEntry


BASE_SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".docx", ".pdf"}
LEGACY_DOC_SUPPORTED = sys.platform == "darwin" and Path("/usr/bin/textutil").exists()
SUPPORTED_EXTENSIONS = BASE_SUPPORTED_EXTENSIONS | ({".doc"} if LEGACY_DOC_SUPPORTED else set())
PDF_REDACTION_SCALE = 2.0
PDF_REDACTION_PADDING = 1.75
OCR_RENDER_SCALE = 2.0
OCR_REDACTION_PADDING = 3
DEFAULT_TESSERACT_LANG = "ita+eng"
OOXML_NAMESPACES = {
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
    "dcmitype": "http://purl.org/dc/dcmitype/",
    "xsi": "http://www.w3.org/2001/XMLSchema-instance",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
    "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
}
TEXT_NODE_NAMES = {"t", "delText", "instrText"}
PERSONAL_ATTRIBUTE_NAMES = {"author", "initials", "lastModifiedBy"}

for prefix, uri in OOXML_NAMESPACES.items():
    ET.register_namespace(prefix, uri)


class OcrUnavailableError(ValueError):
    """Raised when a PDF needs local OCR but Tesseract is not installed."""


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str
    extension: str
    ocr_pages: tuple[int, ...] = ()


@dataclass(frozen=True)
class AnonymizedDocument:
    filename: str
    data: bytes
    text: str
    findings: list[Finding]
    reversible_mapping: tuple[ReversibleMapEntry, ...] = ()


@dataclass(frozen=True)
class PdfRedactionRect:
    left: float
    bottom: float
    right: float
    top: float


@dataclass(frozen=True)
class PdfTextResult:
    text: str
    ocr_pages: tuple[int, ...] = ()


@dataclass(frozen=True)
class OcrWord:
    text: str
    start: int
    end: int
    left: int
    top: int
    right: int
    bottom: int


@dataclass(frozen=True)
class OcrPageText:
    text: str
    words: list[OcrWord]


def load_document(path: str | Path) -> LoadedDocument:
    document_path = Path(path)
    extension = document_path.suffix.lower()
    ocr_pages: tuple[int, ...] = ()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Formato non supportato. Usa uno di questi: {supported}")

    if extension in {".txt", ".md", ".csv"}:
        text = document_path.read_text(encoding="utf-8", errors="replace")
    elif extension == ".doc":
        text = _read_legacy_doc(document_path)
    elif extension == ".docx":
        text = _read_docx(document_path)
    else:
        pdf_text = _read_pdf(document_path)
        text = pdf_text.text
        ocr_pages = pdf_text.ocr_pages

    return LoadedDocument(path=document_path, text=text, extension=extension, ocr_pages=ocr_pages)


def anonymize_loaded_document(
    document: LoadedDocument,
    engine: PrivacyEngine,
    mode: AnonymizationMode = "standard",
    *,
    reversible_entries: Iterable[ReversibleMapEntry] | None = None,
    findings: list[Finding] | None = None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> AnonymizedDocument:
    if mode == "reversible" and document.extension == ".pdf":
        raise ValueError(
            "La modalità reversibile non è disponibile per i PDF: usa Massima protezione per creare un PDF redatto "
            "oppure incolla il testo estratto per lavorare con segnaposti ricostruibili."
        )

    # Il filtro posizionale dei findings pre-calcolati è affidabile solo qui (testo estratto):
    # .docx e .pdf ri-analizzano internamente il testo per parte (nodi XML o pagine), dove
    # l'esclusione avviene per valore esatto tramite excluded_values, in modo fail-closed:
    # un valore che non corrisponde esattamente resta comunque anonimizzato. Le selezioni
    # manuali (extra_values) seguono la stessa logica per valore, in senso opposto: vengono
    # aggiunte come finding ovunque il valore compaia alla lettera in quella parte.
    findings = findings if findings is not None else engine.analyze(document.text, mode)
    findings = _filter_excluded_findings(document.text, findings, excluded_values)
    reversible_session = ReversibleAnonymizer(reversible_entries) if mode == "reversible" else None
    anonymized_text = (
        reversible_session.anonymize(document.text, findings)
        if reversible_session
        else engine.anonymize(document.text, findings, mode)
    )
    output_name = f"{document.path.stem}_anonimizzato{_output_extension(document.extension)}"

    if document.extension in {".txt", ".md", ".csv"}:
        data = anonymized_text.encode("utf-8")
    elif document.extension == ".doc":
        data = _anonymize_legacy_doc(
            document.path,
            engine,
            mode,
            reversible_session=reversible_session,
            excluded_values=excluded_values,
            extra_values=extra_values,
        )
    elif document.extension == ".docx":
        data = _anonymize_docx(
            document.path,
            engine,
            mode,
            reversible_session=reversible_session,
            excluded_values=excluded_values,
            extra_values=extra_values,
        )
    else:
        data = _anonymize_pdf(document.path, engine, mode, excluded_values=excluded_values, extra_values=extra_values)

    return AnonymizedDocument(
        filename=output_name,
        data=data,
        text=anonymized_text,
        findings=findings,
        reversible_mapping=reversible_session.mapping if reversible_session else (),
    )


def _filter_excluded_findings(
    text: str,
    findings: list[Finding],
    excluded_values: frozenset[tuple[str, str]] | None,
) -> list[Finding]:
    """Drop the findings whose (entity_type, exact value) pair was excluded by the user.

    Confronto esatto e case-sensitive, senza normalizzazioni: se il testo ri-analizzato
    (nodo XML, pagina PDF o OCR) non riproduce il valore alla lettera, il finding NON
    viene escluso e il dato resta anonimizzato (fail-closed)."""
    if not excluded_values:
        return findings
    return [
        finding
        for finding in findings
        if (finding.entity_type, text[finding.start : finding.end]) not in excluded_values
    ]


def _add_extra_value_findings(
    text: str,
    findings: list[Finding],
    extra_values: frozenset[tuple[str, str]] | None,
) -> list[Finding]:
    """Aggiunge finding per ogni occorrenza letterale dei valori inclusi manualmente.

    Confronto esatto e case-sensitive, come per le esclusioni; un'occorrenza che si
    sovrappone a un finding già presente viene saltata per non duplicare la redazione."""
    if not extra_values:
        return findings
    result = list(findings)
    occupied = [(finding.start, finding.end) for finding in findings]
    for entity_type, value in extra_values:
        if not value:
            continue
        start = 0
        while True:
            idx = text.find(value, start)
            if idx == -1:
                break
            end = idx + len(value)
            if not any(idx < occupied_end and occupied_start < end for occupied_start, occupied_end in occupied):
                result.append(Finding(entity_type, idx, end, 1.0, source="manual"))
                occupied.append((idx, end))
            start = idx + 1
    return result


def excluded_value_pairs(
    text: str,
    findings: list[Finding],
    included_mask: list[bool],
) -> frozenset[tuple[str, str]]:
    """Compute the (entity_type, value) pairs safe to exclude from a per-occurrence mask.

    Fail-closed: una coppia entra nel risultato solo se OGNI occorrenza di quel valore
    con quel tipo è deselezionata; basta una occorrenza selezionata per mantenerla
    anonimizzata ovunque. Gli indici oltre la maschera contano come selezionati."""
    pair_included: dict[tuple[str, str], bool] = {}
    for index, finding in enumerate(findings):
        pair = (finding.entity_type, text[finding.start : finding.end])
        included = index >= len(included_mask) or included_mask[index]
        pair_included[pair] = pair_included.get(pair, False) or included
    return frozenset(pair for pair, included in pair_included.items() if not included)


def _output_extension(extension: str) -> str:
    if extension == ".pdf":
        return ".pdf"
    if extension in {".doc", ".docx"}:
        return ".docx"
    if extension == ".csv":
        return ".csv"
    return ".txt"


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in _iter_docx_paragraphs(document) if paragraph.text)

    return "\n".join(parts)


def _anonymize_docx(
    path: Path,
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    reversible_session: ReversibleAnonymizer | None = None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> bytes:
    return _sanitize_docx_package(
        path.read_bytes(),
        engine,
        mode,
        reversible_session=reversible_session,
        excluded_values=excluded_values,
        extra_values=extra_values,
    )


def _read_legacy_doc(path: Path) -> str:
    result = _run_textutil("-convert", "txt", "-stdout", str(path))
    return result.stdout.decode("utf-8", errors="replace").strip()


def _anonymize_legacy_doc(
    path: Path,
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    reversible_session: ReversibleAnonymizer | None = None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        converted_path = Path(tmpdir) / f"{path.stem}.docx"
        _run_textutil("-convert", "docx", "-output", str(converted_path), str(path))
        return _anonymize_docx(
            converted_path,
            engine,
            mode,
            reversible_session=reversible_session,
            excluded_values=excluded_values,
            extra_values=extra_values,
        )


def _anonymize_pdf(
    path: Path,
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> bytes:
    from pypdf import PdfReader
    import pypdfium2 as pdfium
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    source_pdf = pdfium.PdfDocument(str(path))
    reader = PdfReader(path)
    output = BytesIO()
    redacted_pdf = canvas.Canvas(output)

    try:
        for page_index in range(len(source_pdf)):
            page = source_pdf[page_index]
            width, height = page.get_size()
            text_page = page.get_textpage()
            bitmap = page.render(scale=PDF_REDACTION_SCALE)
            try:
                image = bitmap.to_pil().convert("RGB")
            finally:
                close = getattr(bitmap, "close", None)
                if close:
                    close()

            page_text = text_page.get_text_range()
            if page_text.strip():
                findings = _filter_excluded_findings(page_text, engine.analyze(page_text, mode), excluded_values)
                findings = _add_extra_value_findings(page_text, findings, extra_values)
                redaction_rects = _pdf_redaction_rects(text_page, page_text, findings, page_index)
                _draw_pdf_redactions(image, (width, height), redaction_rects)

            # A page can contain extractable text and an embedded image. Redacting only
            # the text would leave any personal data in that image visible.
            page_has_images = _pdf_page_has_images(reader.pages[page_index])
            if page_has_images:
                if not _tesseract_available():
                    raise OcrUnavailableError(_ocr_unavailable_message(str(page_index + 1)))
                ocr_text = _ocr_image(image)
                findings = _filter_excluded_findings(ocr_text.text, engine.analyze(ocr_text.text, mode), excluded_values)
                findings = _add_extra_value_findings(ocr_text.text, findings, extra_values)
                _draw_ocr_redactions(image, ocr_text.words, findings)
            elif not page_text.strip():
                if _tesseract_available():
                    ocr_text = _ocr_image(image)
                    findings = _filter_excluded_findings(
                        ocr_text.text, engine.analyze(ocr_text.text, mode), excluded_values
                    )
                    findings = _add_extra_value_findings(ocr_text.text, findings, extra_values)
                    _draw_ocr_redactions(image, ocr_text.words, findings)
                else:
                    raise OcrUnavailableError(_ocr_unavailable_message(str(page_index + 1)))

            redacted_pdf.setPageSize((width, height))
            redacted_pdf.drawImage(ImageReader(image), 0, 0, width=width, height=height)
            redacted_pdf.showPage()
    finally:
        close = getattr(source_pdf, "close", None)
        if close:
            close()

    redacted_pdf.save()
    return output.getvalue()


def _pdf_redaction_rects(
    text_page,
    page_text: str,
    findings: list[Finding],
    page_index: int,
) -> list[PdfRedactionRect]:
    rects: list[PdfRedactionRect] = []
    unmapped_findings: list[str] = []

    for finding in findings:
        finding_rects = _pdf_finding_rects(text_page, page_text, finding)
        if not finding_rects:
            unmapped_findings.append(finding.entity_type)
            continue
        rects.extend(finding_rects)

    if unmapped_findings:
        entities = ", ".join(sorted(set(unmapped_findings)))
        raise ValueError(
            "Non riesco a redigere il PDF in modo sicuro: alcuni dati rilevati non hanno coordinate affidabili "
            f"nella pagina {page_index + 1} ({entities})."
        )

    return rects


def _pdf_finding_rects(text_page, page_text: str, finding: Finding) -> list[PdfRedactionRect]:
    char_boxes = []
    for index in range(finding.start, finding.end):
        if index >= len(page_text) or page_text[index].isspace():
            continue
        try:
            box = text_page.get_charbox(index, loose=True)
        except Exception:
            continue
        if _is_valid_pdf_box(box):
            char_boxes.append(PdfRedactionRect(*box))

    return _merge_pdf_char_boxes(char_boxes)


def _is_valid_pdf_box(box) -> bool:
    if not box or len(box) != 4:
        return False
    left, bottom, right, top = box
    return right > left and top > bottom


def _merge_pdf_char_boxes(boxes: list[PdfRedactionRect]) -> list[PdfRedactionRect]:
    lines: list[list[PdfRedactionRect]] = []

    for box in sorted(boxes, key=lambda item: (-(item.bottom + item.top) / 2, item.left)):
        box_center = (box.bottom + box.top) / 2
        box_height = box.top - box.bottom
        for line in lines:
            line_bottom = min(item.bottom for item in line)
            line_top = max(item.top for item in line)
            line_center = (line_bottom + line_top) / 2
            line_height = line_top - line_bottom
            if abs(box_center - line_center) <= max(box_height, line_height, 1.0) * 0.7:
                line.append(box)
                break
        else:
            lines.append([box])

    return [
        PdfRedactionRect(
            min(box.left for box in line),
            min(box.bottom for box in line),
            max(box.right for box in line),
            max(box.top for box in line),
        )
        for line in lines
    ]


def _draw_pdf_redactions(image, page_size: tuple[float, float], rects: list[PdfRedactionRect]) -> None:
    from PIL import ImageDraw

    _, page_height = page_size
    draw = ImageDraw.Draw(image)
    image_width, image_height = image.size

    for rect in rects:
        left = max(0, int((rect.left - PDF_REDACTION_PADDING) * PDF_REDACTION_SCALE))
        top = max(0, int((page_height - rect.top - PDF_REDACTION_PADDING) * PDF_REDACTION_SCALE))
        right = min(image_width, int((rect.right + PDF_REDACTION_PADDING) * PDF_REDACTION_SCALE) + 1)
        bottom = min(image_height, int((page_height - rect.bottom + PDF_REDACTION_PADDING) * PDF_REDACTION_SCALE) + 1)
        if right > left and bottom > top:
            draw.rectangle((left, top, right, bottom), fill=(0, 0, 0))


def _run_textutil(*args: str) -> subprocess.CompletedProcess[bytes]:
    try:
        return subprocess.run(
            ["/usr/bin/textutil", *args],
            check=True,
            capture_output=True,
        )
    except FileNotFoundError as exc:
        raise RuntimeError("Conversione .doc disponibile solo su macOS con textutil.") from exc
    except subprocess.CalledProcessError as exc:
        detail = exc.stderr.decode("utf-8", errors="replace").strip()
        message = f"Conversione documento non riuscita: {detail}" if detail else "Conversione documento non riuscita."
        raise RuntimeError(message) from exc


def _iter_docx_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _sanitize_docx_package(
    data: bytes,
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    reversible_session: ReversibleAnonymizer | None = None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> bytes:
    source_buffer = BytesIO(data)
    output = BytesIO()

    with zipfile.ZipFile(source_buffer, "r") as source, zipfile.ZipFile(output, "w") as target:
        for item in source.infolist():
            payload = source.read(item.filename)
            if not item.is_dir():
                payload = _sanitize_docx_part(
                    item.filename,
                    payload,
                    engine,
                    mode,
                    reversible_session=reversible_session,
                    excluded_values=excluded_values,
                    extra_values=extra_values,
                )
            target.writestr(item, payload)

    return output.getvalue()


def _sanitize_docx_part(
    name: str,
    data: bytes,
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    reversible_session: ReversibleAnonymizer | None = None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> bytes:
    if _is_docx_metadata_part(name):
        # Lo scrub dei metadati non è filtrabile: i dati personali lì non passano dal pannello.
        return _scrub_metadata_xml(data)
    if _is_docx_text_part(name):
        return _anonymize_ooxml_text(
            data,
            engine,
            mode,
            use_table_context=_is_docx_primary_story_part(name),
            reversible_session=reversible_session,
            excluded_values=excluded_values,
            extra_values=extra_values,
        )
    return data


def _is_docx_metadata_part(name: str) -> bool:
    return name.startswith("docProps/") and name.endswith(".xml")


def _is_docx_text_part(name: str) -> bool:
    if name.startswith("word/") and name.endswith(".xml"):
        filename = Path(name).name
        return (
            filename == "document.xml"
            or filename.startswith("comments")
            or filename.startswith("header")
            or filename.startswith("footer")
            or filename in {"footnotes.xml", "endnotes.xml"}
            or name.startswith("word/glossary/")
        )
    return name.startswith("customXml/") and name.endswith(".xml")


def _is_docx_primary_story_part(name: str) -> bool:
    filename = Path(name).name
    return name.startswith("word/") and (
        filename == "document.xml" or filename.startswith("header") or filename.startswith("footer")
    )


def _scrub_metadata_xml(data: bytes) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data

    for element in root.iter():
        if element.text and element.text.strip():
            element.text = ""
        for attr_name in list(element.attrib):
            local_name = _local_name(attr_name)
            if local_name in PERSONAL_ATTRIBUTE_NAMES:
                element.attrib[attr_name] = ""
            elif local_name in {"name", "linkTarget"}:
                element.attrib[attr_name] = "Anonymized"

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _anonymize_ooxml_text(
    data: bytes,
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    use_table_context: bool,
    reversible_session: ReversibleAnonymizer | None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data

    _strip_personal_ooxml_attributes(root)
    processed_nodes: set[int] = set()

    if use_table_context:
        for table_row in _iter_elements_by_local_name(root, "tr"):
            nodes = list(_iter_text_nodes(table_row))
            _anonymize_text_nodes(
                nodes,
                engine,
                mode,
                reversible_session=reversible_session,
                excluded_values=excluded_values,
                extra_values=extra_values,
            )
            processed_nodes.update(id(node) for node in nodes)

    for container in _iter_text_containers(root):
        nodes = [node for node in _iter_text_nodes(container) if id(node) not in processed_nodes]
        _anonymize_text_nodes(
            nodes,
            engine,
            mode,
            reversible_session=reversible_session,
            excluded_values=excluded_values,
            extra_values=extra_values,
        )
        processed_nodes.update(id(node) for node in nodes)

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _strip_personal_ooxml_attributes(root: ET.Element) -> None:
    for element in root.iter():
        for attr_name in list(element.attrib):
            if _local_name(attr_name) in PERSONAL_ATTRIBUTE_NAMES:
                element.attrib[attr_name] = ""


def _iter_text_containers(root: ET.Element):
    seen: set[int] = set()
    for container in _iter_containers_below(root):
        identity = id(container)
        if identity not in seen:
            seen.add(identity)
            yield container


def _iter_containers_below(root: ET.Element):
    found_container = False
    if _local_name(root.tag) in {"p", "comment", "footnote", "endnote"}:
        found_container = True
        yield root
    for element in root.iter():
        if element is not root and _local_name(element.tag) in {"p", "comment", "footnote", "endnote"}:
            found_container = True
            yield element
    if not found_container and any(True for _ in _iter_text_nodes(root)):
        yield root


def _iter_text_nodes(container: ET.Element):
    for element in container.iter():
        if _local_name(element.tag) in TEXT_NODE_NAMES:
            yield element


def _iter_elements_by_local_name(root: ET.Element, local_name: str):
    for element in root.iter():
        if _local_name(element.tag) == local_name:
            yield element


def _anonymize_text_nodes(
    nodes: list[ET.Element],
    engine: PrivacyEngine,
    mode: AnonymizationMode,
    *,
    reversible_session: ReversibleAnonymizer | None = None,
    excluded_values: frozenset[tuple[str, str]] | None = None,
    extra_values: frozenset[tuple[str, str]] | None = None,
) -> None:
    if not nodes:
        return

    text = "".join(node.text or "" for node in nodes)
    if not text:
        return

    if reversible_session:
        reversible_session.reserve_placeholders(text)

    # Il filtro va applicato PRIMA di assegnare i segnaposti reversibili: un valore
    # escluso non deve mai entrare nella mappa reversibile. Le selezioni manuali (extra_values)
    # vanno invece aggiunte prima dei segnaposti, in modo da entrare correttamente nella mappa.
    findings = _filter_excluded_findings(text, engine.analyze(text, mode), excluded_values)
    findings = _add_extra_value_findings(text, findings, extra_values)
    if not findings:
        return

    replacements = []
    for finding in findings:
        value = text[finding.start : finding.end]
        if reversible_session:
            replacement = reversible_session.placeholder_for(finding.entity_type, value)
        else:
            replacement = engine.anonymize(
                value,
                [Finding(finding.entity_type, 0, finding.end - finding.start, finding.score)],
                mode,
            )
        replacements.append((finding.start, finding.end, replacement))
    replacements.sort(key=lambda replacement: replacement[0])
    cursor = 0

    for node in nodes:
        node_text = node.text or ""
        node_start = cursor
        node_end = node_start + len(node_text)
        cursor = node_end
        pieces: list[str] = []
        text_cursor = node_start

        for start, end, replacement in replacements:
            if end <= node_start or start >= node_end:
                continue
            if start >= text_cursor:
                pieces.append(text[text_cursor:start])
            if node_start <= start < node_end:
                pieces.append(replacement)
            text_cursor = max(text_cursor, min(end, node_end))

        pieces.append(text[text_cursor:node_end])
        node.text = "".join(pieces)


def _local_name(name: str) -> str:
    if "}" in name:
        return name.rsplit("}", 1)[1]
    return name


def _read_pdf(path: Path) -> PdfTextResult:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages: list[str] = []
    ocr_pages: list[int] = []
    unreadable_pages: list[int] = []
    pdfium_doc = None

    try:
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            page_has_images = _pdf_page_has_images(page)
            if text:
                if not page_has_images:
                    pages.append(text)
                    continue
            elif not page_has_images:
                continue

            if not _tesseract_available():
                unreadable_pages.append(page_number)
                continue

            if pdfium_doc is None:
                import pypdfium2 as pdfium

                pdfium_doc = pdfium.PdfDocument(str(path))
            ocr_text = _ocr_pdfium_page(pdfium_doc[page_number - 1])
            if ocr_text.text.strip():
                pages.append(ocr_text.text)
                ocr_pages.append(page_number)
            else:
                unreadable_pages.append(page_number)
    finally:
        close = getattr(pdfium_doc, "close", None)
        if close:
            close()

    if unreadable_pages:
        pages_label = ", ".join(str(page_number) for page_number in unreadable_pages)
        if _tesseract_available():
            raise ValueError(
                "OCR locale completato, ma non ho trovato testo affidabile nelle pagine "
                f"{pages_label}. Controlla la qualità della scansione e riprova."
            )
        raise OcrUnavailableError(_ocr_unavailable_message(pages_label))

    if not pages:
        raise ValueError(
            "Il PDF non contiene testo estraibile. Se è una scansione, abilita OCR locale con Tesseract e riprova."
        )

    return PdfTextResult(text="\n\n".join(pages), ocr_pages=tuple(ocr_pages))


def _ocr_pdfium_page(page) -> OcrPageText:
    bitmap = page.render(scale=OCR_RENDER_SCALE)
    try:
        image = bitmap.to_pil().convert("RGB")
    finally:
        close = getattr(bitmap, "close", None)
        if close:
            close()
    return _ocr_image(image)


def _ocr_image(image) -> OcrPageText:
    command = _tesseract_command()
    if command is None:
        raise OcrUnavailableError(_ocr_unavailable_message(""))

    # mkstemp + reopen by descriptor: NamedTemporaryFile cannot be reopened by name on Windows.
    descriptor, image_path = tempfile.mkstemp(suffix=".png")
    try:
        with os.fdopen(descriptor, "wb") as image_file:
            image.save(image_file, format="PNG")
        return _run_tesseract_tsv(command, image_path)
    finally:
        os.unlink(image_path)


def _run_tesseract_tsv(command: str, image_path: str) -> OcrPageText:
    languages = [os.environ.get("OMISSIS_TESSERACT_LANG", DEFAULT_TESSERACT_LANG)]
    if languages[0] != "eng":
        languages.append("eng")

    last_error = ""
    for language in languages:
        result = subprocess.run(
            [command, image_path, "stdout", "-l", language, "tsv"],
            check=False,
            capture_output=True,
            text=True,
        )
        if result.returncode == 0:
            return _parse_tesseract_tsv(result.stdout)
        last_error = result.stderr.strip()

    detail = f": {last_error}" if last_error else "."
    raise ValueError(f"OCR locale non riuscito{detail}")


def _parse_tesseract_tsv(tsv_text: str) -> OcrPageText:
    words: list[OcrWord] = []
    pieces: list[str] = []
    cursor = 0

    reader = csv.DictReader(StringIO(tsv_text), delimiter="\t")
    for row in reader:
        text = (row.get("text") or "").strip()
        if not text:
            continue
        try:
            left = int(float(row.get("left") or 0))
            top = int(float(row.get("top") or 0))
            width = int(float(row.get("width") or 0))
            height = int(float(row.get("height") or 0))
        except ValueError:
            continue
        if width <= 0 or height <= 0:
            continue

        if pieces:
            pieces.append(" ")
            cursor += 1
        start = cursor
        pieces.append(text)
        cursor += len(text)
        words.append(
            OcrWord(
                text=text,
                start=start,
                end=cursor,
                left=left,
                top=top,
                right=left + width,
                bottom=top + height,
            )
        )

    return OcrPageText(text="".join(pieces), words=words)


def _draw_ocr_redactions(image, words: list[OcrWord], findings: list[Finding]) -> None:
    from PIL import ImageDraw

    draw = ImageDraw.Draw(image)
    image_width, image_height = image.size

    for finding in findings:
        for word in words:
            if word.end <= finding.start or word.start >= finding.end:
                continue
            left = max(0, word.left - OCR_REDACTION_PADDING)
            top = max(0, word.top - OCR_REDACTION_PADDING)
            right = min(image_width, word.right + OCR_REDACTION_PADDING)
            bottom = min(image_height, word.bottom + OCR_REDACTION_PADDING)
            if right > left and bottom > top:
                draw.rectangle((left, top, right, bottom), fill=(0, 0, 0))


def _tesseract_available() -> bool:
    return _tesseract_command() is not None


def _tesseract_command() -> str | None:
    candidates = [
        os.environ.get("OMISSIS_TESSERACT_PATH"),
        shutil.which("tesseract"),
        "/opt/homebrew/bin/tesseract",
        "/usr/local/bin/tesseract",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def _ocr_unavailable_message(pages_label: str) -> str:
    page_detail = f" (pagine: {pages_label})" if pages_label else ""
    return (
        "Il PDF contiene immagini che potrebbero includere dati personali"
        f"{page_detail}. Per anonimizzarlo in sicurezza installa Tesseract OCR locale e riprova: "
        "serve anche quando la pagina contiene testo selezionabile."
    )


def _pdf_page_has_images(page) -> bool:
    resources = _pdf_object(page.get("/Resources"))
    return _pdf_resources_have_images(resources)


def _pdf_resources_have_images(resources, depth: int = 0) -> bool:
    if not resources or depth > 4:
        return False

    xobjects = _pdf_object(resources.get("/XObject")) if hasattr(resources, "get") else None
    if not xobjects:
        return False

    for item in xobjects.values():
        xobject = _pdf_object(item)
        if not hasattr(xobject, "get"):
            continue
        subtype = str(xobject.get("/Subtype"))
        if subtype == "/Image":
            return True
        if subtype == "/Form" and _pdf_resources_have_images(_pdf_object(xobject.get("/Resources")), depth + 1):
            return True

    return False


def _pdf_object(value):
    if hasattr(value, "get_object"):
        return value.get_object()
    return value
